"""
Document Builder - Construction de rapports Word au format CIH Bank
Style procédure bancaire : header table, Généralités, Logigramme, Description des opérations
Couleurs exactes extraites du document de référence CIH Bank.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import base64
import io
import tempfile
import logging
from typing import Dict, List, Any, Optional

logger = logging.getLogger(__name__)

# ── Palette CIH Bank (extraite du document de référence) ─────────────────────
BLUE_H1       = RGBColor(0, 153, 204)    # #0099CC — fond H1 (Généralités, 1. Description)
GREY_H2       = RGBColor(191, 191, 191)  # #BFBFBF — fond H2 (Objet, Périmètre, 1.1, 1.3...)
GREY_CELL     = RGBColor(242, 242, 242)  # #F2F2F2 — fond cellule label gauche tableau opération
WHITE         = RGBColor(255, 255, 255)
BLACK         = RGBColor(0, 0, 0)
GREY_TEXT     = RGBColor(89, 89, 89)
BORDER_LIGHT  = 'E7E6E6'                 # bordures tableau header
BORDER_TABLE  = 'A6A6A6'                 # bordures tableaux opérations


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Supprimer ancien shd si présent
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_para_bg(p, hex_color: str):
    """Fond coloré sur un paragraphe (pour H1/H2)."""
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:shd')):
        pPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def _set_table_borders(table, color=BORDER_TABLE, sz=4):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_col_width(table, col_idx: int, width_cm: float):
    for row in table.rows:
        cell = row.cells[col_idx]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is None:
            tcW = OxmlElement('w:tcW')
            tcPr.append(tcW)
        tcW.set(qn('w:w'), str(int(width_cm * 567)))
        tcW.set(qn('w:type'), 'dxa')


def _run(paragraph, text, bold=False, size=10, color=None, italic=False, name='Calibri'):
    r = paragraph.add_run(text)
    r.font.name = name
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return r


class DocumentBuilder:
    """
    Génère des documents Word au format procédure CIH Bank.
    Structure : Header table → Généralités → Sommaire → 1. Description procédure
    """

    def __init__(self):
        self.doc = Document()
        self.step_id_to_name: Dict[str, str] = {}
        self._set_margins()
        self._apply_cih_styles()

    def _set_margins(self):
        for section in self.doc.sections:
            section.top_margin    = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

    def _apply_cih_styles(self):
        styles = self.doc.styles

        normal = styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(10)
        normal.font.color.rgb = BLACK
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal.paragraph_format.line_spacing = 1.15
        normal.paragraph_format.space_after = Pt(4)

        # H1 — fond #0099CC, texte blanc (style TOC)
        h1 = styles['Heading 1']
        h1.font.name = 'Calibri'
        h1.font.size = Pt(11)
        h1.font.bold = True
        h1.font.color.rgb = WHITE
        h1.paragraph_format.space_before = Pt(10)
        h1.paragraph_format.space_after = Pt(4)
        h1.paragraph_format.keep_with_next = True

        # H2 — fond #BFBFBF, texte blanc
        h2 = styles['Heading 2']
        h2.font.name = 'Calibri'
        h2.font.size = Pt(10)
        h2.font.bold = True
        h2.font.color.rgb = WHITE
        h2.paragraph_format.space_before = Pt(8)
        h2.paragraph_format.space_after = Pt(4)
        h2.paragraph_format.keep_with_next = True

        try:
            h3 = styles['Heading 3']
        except KeyError:
            h3 = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        h3.font.name = 'Calibri'
        h3.font.size = Pt(10)
        h3.font.bold = True
        h3.font.color.rgb = GREY_TEXT
        h3.paragraph_format.space_before = Pt(8)
        h3.paragraph_format.space_after = Pt(4)

        logger.info("✅ Styles CIH Bank appliqués")

    # ── Point d'entrée ────────────────────────────────────────────────────────

    def generate_process_report(
        self,
        metadata: Dict[str, Any],
        workflow: List[Dict[str, str]],
        enrichments: Dict[str, Dict[str, str]],
        diagram_image: Optional[str] = None,
        options: Dict[str, Any] = None
    ) -> str:
        if options is None:
            options = {'include_diagram': True, 'include_enrichments': True, 'detail_level': 'standard'}

        logger.info(f"🔨 Génération procédure '{metadata.get('nom', 'Procédure')}'")

        self.step_id_to_name = {s['id']: s['étape'] for s in workflow}

        self._add_procedure_header(metadata)
        self._add_generalites(metadata)
        self._add_table_of_contents()
        self._add_h1("1. Description de la procédure")
        self._add_regles_gestion(metadata)

        if options.get('include_diagram') and diagram_image:
            self._add_logigramme(diagram_image)
        else:
            self._add_logigramme_placeholder()

        self._add_description_operations(workflow, enrichments, options)

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx', prefix='processmate_')
        self.doc.save(tmp.name)
        logger.info(f"💾 Document sauvegardé : {tmp.name}")
        return tmp.name

    # ── 1. Header ────────────────────────────────────────────────────────────

    def _add_procedure_header(self, metadata: Dict[str, Any]):
        """
        Pôle/Direction centré au-dessus, puis tableau 2 colonnes :
        gauche = Procédure (gris) + titre (noir gras) | droite = Réf/Version/Dates
        Bordures légères #E7E6E6.
        """
        # Pôle / Direction au-dessus (centré, petite police)
        p_pole = self.doc.add_paragraph()
        p_pole.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p_pole, metadata.get('pole', "Pôle Systèmes d'information"), size=10, color=BLACK)

        p_dir = self.doc.add_paragraph()
        p_dir.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p_dir, metadata.get('direction', 'Direction Organisation et Reengineering de Processus'), size=10, color=BLACK)

        self.doc.add_paragraph()  # espace

        # Tableau header
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        _set_table_borders(table, color=BORDER_LIGHT, sz=4)

        cell_left  = table.cell(0, 0)
        cell_right = table.cell(0, 1)

        # Colonne gauche : "Procédure" en gris/cyan + titre en gras noir
        p_proc = cell_left.paragraphs[0]
        p_proc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p_proc, "Procédure", bold=True, size=16, color=RGBColor(166, 166, 166))

        p_nom = cell_left.add_paragraph()
        p_nom.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p_nom, metadata.get('nom', 'Procédure'), bold=True, size=16, color=BLACK)

        # Colonne droite : Réf / Version / Dates
        p_ref = cell_right.paragraphs[0]
        _run(p_ref, f"Réf : {metadata.get('ref', '')}", size=9)

        for label, key in [
            ("Version : ", 'version'),
            ("Date de prise d'effet : ", 'dateEffet'),
            ("Date diffusion : ", 'dateDiffusion'),
        ]:
            p = cell_right.add_paragraph()
            _run(p, f"{label}{metadata.get(key, '')}", size=9)

        # Largeurs : 60% / 40%
        total_cm = 16.0
        _set_col_width(table, 0, total_cm * 0.60)
        _set_col_width(table, 1, total_cm * 0.40)

        self.doc.add_paragraph()

    # ── 2. Généralités ───────────────────────────────────────────────────────

    def _add_generalites(self, metadata: Dict[str, Any]):
        self._add_h1("Généralités")

        self._add_h2("Objet")
        objet = metadata.get('objet', '')
        if objet:
            self.doc.add_paragraph(objet)

        self._add_h2("Périmètre d'application")
        perimeter = metadata.get('perimeter', '')
        if perimeter:
            self.doc.add_paragraph(perimeter)

        self._add_h2("Responsabilités")
        internes = metadata.get('responsabilites_internes', [])
        externes = metadata.get('responsabilites_externes', [])
        if internes:
            p = self.doc.add_paragraph()
            _run(p, "Acteurs internes :", bold=True)
            for actor in internes:
                bp = self.doc.add_paragraph(style='List Bullet')
                _run(bp, actor)
        if externes:
            p = self.doc.add_paragraph()
            _run(p, "Acteurs externes :", bold=True)
            for actor in externes:
                bp = self.doc.add_paragraph(style='List Bullet')
                _run(bp, actor)
        if not internes and not externes:
            self.doc.add_paragraph()

        self._add_h2("Références")
        references = metadata.get('references', '')
        if references:
            if isinstance(references, list):
                for ref in references:
                    bp = self.doc.add_paragraph(style='List Bullet')
                    _run(bp, ref)
            else:
                self.doc.add_paragraph(references)
        else:
            self.doc.add_paragraph()

        self._add_h2("Définitions")
        definitions = metadata.get('definitions', [])
        if definitions:
            for item in definitions:
                if isinstance(item, dict):
                    p = self.doc.add_paragraph()
                    _run(p, f"{item.get('terme', '')} : ", bold=True)
                    _run(p, item.get('definition', ''))
        else:
            self.doc.add_paragraph()

        self._add_h2("Abréviations")
        abbreviations = metadata.get('abbreviations', [])
        if abbreviations:
            for item in abbreviations:
                if isinstance(item, dict):
                    p = self.doc.add_paragraph()
                    _run(p, f"{item.get('abrv', '')} : ", bold=True)
                    _run(p, item.get('signification', ''))
        else:
            self.doc.add_paragraph()

        self.doc.add_page_break()

    # ── 3. Sommaire ──────────────────────────────────────────────────────────

    def _add_table_of_contents(self):
        p_title = self.doc.add_paragraph()
        _run(p_title, "Sommaire :", bold=True, size=11, color=BLACK)

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)

        note = self.doc.add_paragraph()
        _run(note, "(Clic droit → « Mettre à jour les champs » dans Word pour générer la table)",
             italic=True, size=8, color=GREY_TEXT)

        self.doc.add_page_break()

    # ── 4. Règles de gestion ─────────────────────────────────────────────────

    def _add_regles_gestion(self, metadata: Dict[str, Any]):
        self._add_h2("1.1 Règles de gestion")
        regles = metadata.get('regles_gestion', '')
        if regles:
            if isinstance(regles, list):
                for regle in regles:
                    bp = self.doc.add_paragraph(style='List Bullet')
                    _run(bp, regle)
            else:
                for line in regles.split('\n'):
                    line = line.strip()
                    if line:
                        self.doc.add_paragraph(line)
        else:
            self.doc.add_paragraph()

    # ── 5. Logigramme ────────────────────────────────────────────────────────

    def _add_logigramme(self, diagram_base64: str):
        self._add_h2("1.2 Logigramme")
        try:
            b64 = diagram_base64.split(',')[1] if diagram_base64.startswith('data:image') else diagram_base64
            image_data = base64.b64decode(b64)
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(image_data), width=Inches(6.3))
            logger.info(f"✅ Logigramme inséré ({len(image_data)} bytes)")
        except Exception as e:
            logger.error(f"❌ Erreur logigramme : {e}")
            p = self.doc.add_paragraph()
            _run(p, f"[Logigramme non disponible : {e}]", color=RGBColor(200, 0, 0))
        self.doc.add_page_break()

    def _add_logigramme_placeholder(self):
        self._add_h2("1.2 Logigramme")
        p = self.doc.add_paragraph()
        _run(p, "[Logigramme non inclus]", italic=True, color=GREY_TEXT)
        self.doc.add_page_break()

    # ── 6. Description des opérations ────────────────────────────────────────

    def _add_description_operations(self, workflow, enrichments, options):
        self._add_h2("1.3 Description des opérations")

        op_number = 0
        for step in workflow:
            op_number += 1
            enrichment = enrichments.get(step['id']) if options.get('include_enrichments') else None
            self._add_operation_block(step, enrichment, op_number)

        logger.info(f"✅ {op_number} opérations documentées")

    def _add_operation_block(self, step, enrichment, number):
        """
        Format CIH Bank :
        Numéro. Titre (gras)
        Elément déclencheur : ... (italique, si présent)
        Tableau 2 colonnes : Acteur | valeur / Description | valeur
        Applicatif : xxx (si présent)
        """
        # Titre
        p_title = self.doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(8)
        _run(p_title, f"{number}.  {step['étape']}", bold=True)

        # Élément déclencheur
        declencheur = ''
        if step['typeBpmn'] == 'StartEvent':
            declencheur = (enrichment.get('declencheur', '') if enrichment else '') or step.get('condition', '')
        elif enrichment:
            declencheur = enrichment.get('declencheur', '')

        if declencheur:
            p_dec = self.doc.add_paragraph()
            _run(p_dec, f"Elément déclencheur : {declencheur}", italic=True, size=9)

        # Tableau Acteur / Description
        table = self.doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        _set_table_borders(table, color=BORDER_TABLE, sz=4)

        # Ligne 1 — Acteur
        c_lbl_a = table.cell(0, 0)
        c_val_a = table.cell(0, 1)
        _set_cell_bg(c_lbl_a, 'F2F2F2')

        p_al = c_lbl_a.paragraphs[0]
        _run(p_al, "Acteur :", bold=True)

        acteur = step.get('acteur', '')
        dept   = step.get('département', '')
        p_av = c_val_a.paragraphs[0]
        _run(p_av, f"{acteur} - {dept}" if dept else acteur)

        # Ligne 2 — Description / Commentaire
        c_lbl_d = table.cell(1, 0)
        c_val_d = table.cell(1, 1)
        _set_cell_bg(c_lbl_d, 'F2F2F2')
        c_lbl_d.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        c_val_d.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        label_desc = "Commentaire :" if step['typeBpmn'] in ['ExclusiveGateway', 'StartEvent', 'EndEvent'] else "Description :"
        _run(c_lbl_d.paragraphs[0], label_desc, bold=True)

        desc_lines = self._build_description(step, enrichment)
        first = True
        for line in desc_lines:
            p = c_val_d.paragraphs[0] if first else c_val_d.add_paragraph()
            first = False
            if line.startswith('• '):
                try:
                    p.style = self.doc.styles['List Bullet']
                except Exception:
                    pass
                _run(p, line[2:])
            else:
                _run(p, line)

        # Largeurs 22% / 78%
        total_cm = 16.0
        _set_col_width(table, 0, total_cm * 0.22)
        _set_col_width(table, 1, total_cm * 0.78)

        # Applicatif
        outil = step.get('outil', '').strip() or (enrichment.get('applicatif', '').strip() if enrichment else '')
        if outil:
            p_app = self.doc.add_paragraph()
            _run(p_app, "Applicatif : ", bold=True, size=9)
            _run(p_app, outil, size=9)

        self.doc.add_paragraph()

    def _build_description(self, step, enrichment):
        lines = []
        type_bpmn = step['typeBpmn']

        if enrichment and enrichment.get('descriptif'):
            for line in enrichment['descriptif'].split('\n'):
                line = line.strip()
                if line:
                    lines.append(line)

        if type_bpmn == 'ExclusiveGateway':
            condition = step.get('condition', '')
            if condition and not lines:
                lines.append(f"Point de décision : {condition}")
            for out in step.get('outputs', []):
                label     = out.get('label', '')     if isinstance(out, dict) else out.label
                target_id = out.get('targetId', '')  if isinstance(out, dict) else out.targetId
                target_name = self.step_id_to_name.get(target_id, target_id)
                if label and target_name:
                    lines.append(f"• Si {label} → {target_name}")
        elif type_bpmn == 'EndEvent' and not lines:
            lines.append("Fin du processus.")
        elif type_bpmn == 'StartEvent' and not lines:
            lines.append("Début du processus.")

        if enrichment:
            if enrichment.get('duree_estimee'):
                lines.append(f"Durée estimée : {enrichment['duree_estimee']}")
            if enrichment.get('frequence'):
                lines.append(f"Fréquence : {enrichment['frequence']}")
            if enrichment.get('kpi'):
                lines.append(f"KPI : {enrichment['kpi']}")

        return lines if lines else ['—']

    # ── Helpers titres ────────────────────────────────────────────────────────

    def _add_h1(self, text: str):
        """H1 avec fond #0099CC et texte blanc."""
        p = self.doc.add_paragraph(style='Heading 1')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_para_bg(p, '0099CC')
        _run(p, text, bold=True, size=11, color=WHITE)
        return p

    def _add_h2(self, text: str):
        """H2 avec fond #BFBFBF et texte blanc."""
        p = self.doc.add_paragraph(style='Heading 2')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_para_bg(p, 'BFBFBF')
        _run(p, text, bold=True, size=10, color=WHITE)
        return p