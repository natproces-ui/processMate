"""
Corrections & Révisions — ProcessMate
Upload d'un PDF annoté → Gemini détecte les remarques → affichage 3 colonnes
"""
from __future__ import annotations
import io
import logging
import uuid
from datetime import datetime
from typing import Optional

from fastapi import APIRouter, File, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from processor.corrections_processor import analyze_pdf_corrections

# python-docx pour l'export
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/corrections", tags=["Corrections"])

# ─── Session store en mémoire ─────────────────────────────────
# (mêmes pattern que process_discovery)
_sessions: dict[str, dict] = {}

MAX_FILE_SIZE = 30 * 1024 * 1024  # 30 Mo


# ─── Modèles ──────────────────────────────────────────────────

class RemarkStatusUpdate(BaseModel):
    status: str   # "pending" | "treated" | "ignored"


# ─── Endpoints ────────────────────────────────────────────────

@router.post("/analyze")
async def analyze_document(file: UploadFile = File(...)):
    """Upload un PDF et retourne les remarques détectées par Gemini."""
    if not file.content_type or "pdf" not in file.content_type.lower():
        raise HTTPException(400, "Seuls les fichiers PDF sont acceptés")

    pdf_bytes = await file.read()
    if len(pdf_bytes) > MAX_FILE_SIZE:
        raise HTTPException(413, f"Fichier trop volumineux (max {MAX_FILE_SIZE // (1024*1024)} Mo)")
    if len(pdf_bytes) == 0:
        raise HTTPException(400, "Le fichier est vide")

    try:
        result = await analyze_pdf_corrections(pdf_bytes, file.filename or "document.pdf")
    except ValueError as e:
        raise HTTPException(422, str(e))
    except Exception as e:
        logger.error("Erreur analyse corrections: %s", e, exc_info=True)
        raise HTTPException(500, f"Erreur lors de l'analyse IA : {e}")

    _sessions[result["session_id"]] = result
    return {"success": True, **result}


@router.get("/sessions/{session_id}")
async def get_session(session_id: str):
    """Récupère une session d'analyse (remarques + synthèse)."""
    session = _sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session introuvable ou expirée")
    return {"success": True, **session}


@router.patch("/sessions/{session_id}/remarks/{remark_id}")
async def update_remark_status(session_id: str, remark_id: str, body: RemarkStatusUpdate):
    """Met à jour le statut d'une remarque (pending / treated / ignored)."""
    session = _sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session introuvable")
    if body.status not in ("pending", "treated", "ignored"):
        raise HTTPException(400, "Statut invalide")

    for remark in session.get("remarks", []):
        if remark["id"] == remark_id:
            remark["status"] = body.status
            return {"success": True}

    raise HTTPException(404, "Remarque introuvable")


@router.get("/sessions/{session_id}/report")
async def download_corrections_report(session_id: str):
    """Génère et retourne le rapport Word des corrections."""
    session = _sessions.get(session_id)
    if not session:
        raise HTTPException(404, "Session introuvable")

    try:
        docx_bytes = _build_corrections_docx(session)
    except Exception as e:
        logger.error("Erreur génération rapport corrections: %s", e)
        raise HTTPException(500, "Erreur lors de la génération du rapport")

    title_slug = session.get("document_title", "rapport").replace(" ", "_")[:40]
    filename = f"corrections_{title_slug}.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─── Génération rapport Word ───────────────────────────────────

BLUE   = RGBColor(0x00, 0x27, 0x6D)
BLUE_L = RGBColor(0x1A, 0x6B, 0xC4)
PALE   = RGBColor(0xD9, 0xE8, 0xF7)
GREY   = RGBColor(0xF5, 0xF5, 0xF5)
GTXT   = RGBColor(0x40, 0x40, 0x40)
GREEN  = RGBColor(0x1B, 0x87, 0x3A)
ORANGE = RGBColor(0xE0, 0x7B, 0x00)
RED    = RGBColor(0xC0, 0x2B, 0x2B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

CRIT_COLORS = {"haute": RED, "moyenne": ORANGE, "faible": GREEN}
STATUS_LABELS = {"pending": "En attente", "treated": "Traité", "ignored": "Ignoré"}
STATUS_COLORS = {"pending": ORANGE, "treated": GREEN, "ignored": GTXT}
TYPE_LABELS = {
    "surlignement": "Surlignement",
    "manuscrit":    "Manuscrit",
    "rature":       "Rature",
    "soulignement": "Soulignement",
    "encadrement":  "Encadrement",
    "diagramme":    "Annotation diagramme",
    "commentaire":  "Commentaire",
}


def _hex(c: RGBColor) -> str:
    return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"

def _cell_bg(cell, color: RGBColor):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    for o in pr.findall(qn("w:shd")): pr.remove(o)
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), _hex(color)); pr.append(s)

def _cell_border(cell):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "left", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "D0D0D0"); borders.append(el)
    pr.append(borders)

def _run(para, text, bold=False, italic=False, size=10, color=None, font="Calibri"):
    r = para.add_run(text)
    r.font.bold = bold; r.font.italic = italic
    r.font.size = Pt(size); r.font.name = font
    if color: r.font.color.rgb = color
    return r

def _spacing(para, before=0, after=4):
    pf = para.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)

def _heading1(doc, text):
    p = doc.add_paragraph()
    _spacing(p, before=14, after=6)
    pPr = p._p.get_or_add_pPr()
    s = OxmlElement("w:shd"); s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto"); s.set(qn("w:fill"), _hex(BLUE)); pPr.append(s)
    _run(p, f"  {text}", bold=True, size=12, color=WHITE)


def _build_corrections_docx(session: dict) -> bytes:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3);  sec.right_margin = Cm(2.5)

    remarks = session.get("remarks", [])
    now_str = datetime.utcnow().strftime("%d/%m/%Y à %H:%M")
    total = len(remarks)
    treated = sum(1 for r in remarks if r.get("status") == "treated")
    pending = sum(1 for r in remarks if r.get("status") == "pending")

    # Page de garde
    cover = doc.add_paragraph()
    _spacing(cover, before=30, after=8)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = cover._p.get_or_add_pPr()
    s = OxmlElement("w:shd"); s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto"); s.set(qn("w:fill"), _hex(BLUE)); pPr.append(s)
    _run(cover, "  RAPPORT DE RÉVISION — CORRECTIONS  ", bold=True, size=16, color=WHITE)

    t = doc.add_paragraph(session.get("document_title", "Document"))
    _spacing(t, before=10, after=6); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(t, "", bold=True, size=14, color=BLUE)
    t.runs[0].text = session.get("document_title", "Document")

    # Stats couverture
    tbl = doc.add_table(rows=2, cols=3); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for j, h in enumerate(["Total remarques", "Traitées", "En attente"]):
        c = tbl.rows[0].cells[j]; c.text = h
        _cell_bg(c, BLUE); _cell_border(c)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].font.bold = True
            c.paragraphs[0].runs[0].font.color.rgb = WHITE
            c.paragraphs[0].runs[0].font.size = Pt(9)
    for j, v in enumerate([str(total), str(treated), str(pending)]):
        c = tbl.rows[1].cells[j]; c.text = v
        _cell_bg(c, PALE); _cell_border(c)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].font.bold = True
            c.paragraphs[0].runs[0].font.color.rgb = BLUE
            c.paragraphs[0].runs[0].font.size = Pt(14)

    gd = doc.add_paragraph(f"\nGénéré le {now_str}")
    _spacing(gd, before=8); gd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if gd.runs: gd.runs[0].font.size = Pt(9); gd.runs[0].font.color.rgb = GTXT

    doc.add_page_break()

    # Synthèse globale
    _heading1(doc, "1. Synthèse globale")
    synth = doc.add_paragraph(session.get("synthese") or "Aucune synthèse disponible.")
    _spacing(synth, before=4, after=8)
    if synth.runs:
        synth.runs[0].font.size = Pt(10); synth.runs[0].font.color.rgb = GTXT

    # Tableau des remarques
    _heading1(doc, "2. Détail des remarques")

    if not remarks:
        p = doc.add_paragraph("Aucune remarque détectée dans ce document.")
        _spacing(p, before=4)
        if p.runs: p.runs[0].font.size = Pt(10); p.runs[0].font.color.rgb = GTXT
    else:
        tbl2 = doc.add_table(rows=1, cols=6)
        tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl2.style = "Table Grid"
        widths = [Inches(0.3), Inches(0.5), Inches(1.0), Inches(1.5), Inches(2.2), Inches(1.2)]
        hdrs = ["#", "Page", "Type", "Texte concerné", "Suggestion de correction", "Statut"]
        for j, (h, w) in enumerate(zip(hdrs, widths)):
            c = tbl2.rows[0].cells[j]
            c.width = w; c.text = h
            _cell_bg(c, BLUE); _cell_border(c)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.bold = True
                c.paragraphs[0].runs[0].font.color.rgb = WHITE
                c.paragraphs[0].runs[0].font.size = Pt(8)

        for idx, r in enumerate(sorted(remarks, key=lambda x: x.get("page", 1)), 1):
            row = tbl2.add_row()
            status = r.get("status", "pending")
            crit = r.get("criticite", "moyenne")
            bg = GREY if idx % 2 == 0 else WHITE
            vals = [
                str(idx),
                str(r.get("page", "")),
                TYPE_LABELS.get(r.get("type", ""), r.get("type", "")),
                r.get("texte_concerne") or r.get("zone") or "—",
                r.get("suggestion") or "—",
                STATUS_LABELS.get(status, status),
            ]
            for j, (val, w) in enumerate(zip(vals, widths)):
                c = row.cells[j]
                c.width = w; c.text = val
                _cell_bg(c, bg); _cell_border(c)
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if j > 2 else WD_ALIGN_PARAGRAPH.CENTER
                if c.paragraphs[0].runs:
                    rn = c.paragraphs[0].runs[0]
                    rn.font.size = Pt(8); rn.font.name = "Calibri"
                    if j == 0:
                        rn.font.color.rgb = CRIT_COLORS.get(crit, ORANGE)
                        rn.font.bold = True
                    elif j == 5:
                        rn.font.color.rgb = STATUS_COLORS.get(status, GTXT)
                        rn.font.bold = True
                    else:
                        rn.font.color.rgb = GTXT

    # Pied de page
    footer = doc.add_paragraph()
    _spacing(footer, before=20)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(footer, f"Document généré par ProcessMate · {now_str}", italic=True, size=8, color=GTXT)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
