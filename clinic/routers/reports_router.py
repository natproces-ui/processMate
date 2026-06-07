"""
Génération de rapports Word pour les campagnes de formalisation.
GET /api/campaigns/{id}/report → téléchargement .docx
"""
from __future__ import annotations
import io
import logging
from datetime import datetime

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse

from database.supabase_client import get_supabase
from processor.campaign_report_processor import generate_campaign_narrative

# python-docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/campaigns", tags=["Rapports"])

# ─── Palette couleurs ──────────────────────────────────────────

BLUE_PRIMARY  = RGBColor(0x00, 0x27, 0x6D)   # #00276D bleu foncé
BLUE_LIGHT    = RGBColor(0x1A, 0x6B, 0xC4)   # #1A6BC4 bleu moyen
BLUE_PALE     = RGBColor(0xD9, 0xE8, 0xF7)   # #D9E8F7 bleu très clair
GREY_LIGHT    = RGBColor(0xF5, 0xF5, 0xF5)   # #F5F5F5
GREY_TEXT     = RGBColor(0x40, 0x40, 0x40)   # #404040
GREEN         = RGBColor(0x1B, 0x87, 0x3A)   # #1B873A
ORANGE        = RGBColor(0xE0, 0x7B, 0x00)   # #E07B00
RED           = RGBColor(0xC0, 0x2B, 0x2B)   # #C02B2B
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)

DIAGNOSTIC_COLOR = {"positif": GREEN, "attention": ORANGE, "critique": RED}
DIAGNOSTIC_LABEL = {"positif": "POSITIF", "attention": "ATTENTION REQUISE", "critique": "CRITIQUE"}

PROC_STATUS_FR = {
    "pending":    "En attente",
    "in_progress":"En cours",
    "formalized": "Formalisée",
    "validated":  "Validée",
    "skipped":    "Ignorée",
}
PROJECT_STATUS_FR = {
    "draft":     "Brouillon",
    "active":    "En cours",
    "completed": "Terminé",
    "archived":  "Archivé",
    "blocked":   "Bloqué",
    "on_hold":   "En pause",
}
CAMPAIGN_STATUS_FR = PROJECT_STATUS_FR  # alias rétrocompat


# ─── Helpers Word ──────────────────────────────────────────────

def _hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"

def _set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(color))
    tcPr.append(shd)

def _set_cell_border(cell, sides=("top", "bottom", "left", "right"), color="D0D0D0", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)

def _para_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY

def _heading(doc: Document, text: str, level: int = 1):
    """Titre de section numéroté avec fond bleu (level=1) ou ligne bleue (level=2)."""
    if level == 1:
        para = doc.add_paragraph()
        _para_spacing(para, before=14, after=6)
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), _hex(BLUE_PRIMARY))
        pPr.append(shd)
        run = para.add_run(f"  {text}")
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = WHITE
        run.font.name = "Calibri"
        # Applique le style Heading 1 pour le sommaire
        para.style = doc.styles["Heading 1"]
        # Réapplique la mise en forme (le style peut écraser)
        run.font.bold = True; run.font.size = Pt(12)
        run.font.color.rgb = WHITE; run.font.name = "Calibri"
    else:
        para = doc.add_paragraph()
        _para_spacing(para, before=8, after=4)
        run = para.add_run(text)
        run.font.bold = True; run.font.size = Pt(11)
        run.font.color.rgb = BLUE_LIGHT; run.font.name = "Calibri"
        para.style = doc.styles["Heading 2"]
        run.font.bold = True; run.font.size = Pt(11)
        run.font.color.rgb = BLUE_LIGHT; run.font.name = "Calibri"

def _body(doc: Document, text: str):
    para = doc.add_paragraph(text)
    _para_spacing(para, before=2, after=4, line=14)
    para.runs[0].font.size = Pt(10)
    para.runs[0].font.color.rgb = GREY_TEXT
    para.runs[0].font.name = "Calibri"
    return para

def _add_toc(doc: Document):
    """Insère un champ TOC Word natif qui se génère automatiquement à l'ouverture."""
    para = doc.add_paragraph()
    _para_spacing(para, before=0, after=12)
    run = para.add_run()
    r = run._r

    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Sommaire"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")

    r.append(begin); r.append(instr); r.append(sep)
    r2 = OxmlElement("w:r"); r2.append(placeholder); para._p.append(r2)
    r.append(end)


def _enable_auto_update_fields(doc: Document):
    """Force Word/LibreOffice à mettre à jour le TOC et tous les champs à l'ouverture."""
    settings = doc.settings.element
    # Évite les doublons
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    existing = settings.find(f"{{{ns}}}updateFields")
    if existing is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)


# ─── Helpers campagne ──────────────────────────────────────────

import uuid as _uuid_mod

def _is_uuid(val: str) -> bool:
    try:
        _uuid_mod.UUID(str(val))
        return True
    except Exception:
        return False


def _get_campaign_full(campaign_id: str) -> dict:
    db = get_supabase()
    res = db.table("formalization_campaigns").select("*").eq("id", campaign_id).execute()
    if not res.data:
        raise HTTPException(404, "Campagne introuvable")
    campaign = res.data[0]

    cp_res = db.table("campaign_procedures").select("*").eq("campaign_id", campaign_id).execute()
    procs = cp_res.data or []

    if procs:
        pids = [p["procedure_id"] for p in procs]

        # ── Lifecycle depuis workflows ──────────────────────────
        wf_res = db.table("workflows").select("id, procedure_metadata_json").in_("id", pids).execute()
        wf_map = {r["id"]: r for r in (wf_res.data or [])}

        # ── Responsables depuis procedure_assignments (RACI) ────
        # Priorité : raci_role=R > assignment_type=owner > autres
        asgn_res = db.table("procedure_assignments").select(
            "procedure_id, raci_role, assignment_type, "
            "user_profiles!procedure_assignments_user_id_fkey(id, full_name, display_name)"
        ).in_("procedure_id", pids).execute()

        responsible_map: dict[str, str] = {}
        for row in (asgn_res.data or []):
            pid = row.get("procedure_id", "")
            profile = row.get("user_profiles") or {}
            name = (profile.get("full_name") or profile.get("display_name") or "").strip()
            if not name:
                continue
            raci  = row.get("raci_role", "") or ""
            atype = row.get("assignment_type", "") or ""
            current = responsible_map.get(pid)
            if current is None:
                responsible_map[pid] = name
            elif raci == "R":
                responsible_map[pid] = name          # R écrase tout
            elif atype == "owner" and current and raci not in ("R",):
                responsible_map[pid] = name           # owner écrase les autres rôles

        # ── Résolution des assigned_to UUID (campaign_procedures) ─
        at_uuids = [p["assigned_to"] for p in procs if p.get("assigned_to") and _is_uuid(str(p["assigned_to"]))]
        user_name_map: dict[str, str] = {}
        if at_uuids:
            up_res = db.table("user_profiles").select("id, full_name, display_name").in_("id", at_uuids).execute()
            for u in (up_res.data or []):
                user_name_map[u["id"]] = (u.get("full_name") or u.get("display_name") or "").strip()

        # ── Enrichissement de chaque procédure ──────────────────
        for p in procs:
            pid = p["procedure_id"]
            wf = wf_map.get(pid)
            if wf:
                meta = wf.get("procedure_metadata_json") or {}
                stages = meta.get("lifecycle_stages") or []
                n_done = sum(1 for s in stages if s.get("status") == "completed")
                p["lifecycle"] = {"stages_done": n_done, "stages_total": len(stages) or 6}
                if not p.get("procedure_nom"):
                    p["procedure_nom"] = meta.get("nom", "")
                if not p.get("procedure_ref"):
                    p["procedure_ref"] = meta.get("ref", "")

            # Résolution responsable : RACI → assigned_to résolu → valeur brute → —
            if pid in responsible_map:
                p["responsible_name"] = responsible_map[pid]
            elif p.get("assigned_to"):
                at = str(p["assigned_to"])
                p["responsible_name"] = user_name_map.get(at) or (at if not _is_uuid(at) else "—")
            else:
                p["responsible_name"] = "—"

    total = len(procs)
    done = sum(1 for p in procs if p.get("status") in ("formalized", "validated"))
    in_progress = sum(1 for p in procs if p.get("status") == "in_progress")
    pending = total - done - in_progress

    campaign["procedures"] = procs
    campaign["stats"] = {
        "total": total, "done": done,
        "in_progress": in_progress, "pending": pending,
        "progress_pct": round((done / total * 100) if total else 0),
    }
    return campaign


# ─── Assemblage Word ───────────────────────────────────────────

def _build_docx(campaign: dict, narrative: dict) -> bytes:
    doc = Document()

    # ── Marges ──────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # Force la mise à jour automatique du TOC à l'ouverture
    _enable_auto_update_fields(doc)

    stats = campaign.get("stats", {})
    procs = campaign.get("procedures", [])
    now_str = datetime.utcnow().strftime("%d/%m/%Y à %H:%M")
    diagnostic = narrative.get("diagnostic", "attention")

    # ── PAGE DE GARDE ────────────────────────────────────────────
    # Bandeau titre
    cover = doc.add_paragraph()
    _para_spacing(cover, before=40, after=8)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = cover._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), _hex(BLUE_PRIMARY))
    pPr.append(shd)
    r = cover.add_run("  RAPPORT DE PROJET DE FORMALISATION  ")
    r.font.bold = True; r.font.size = Pt(18)
    r.font.color.rgb = WHITE; r.font.name = "Calibri"

    # Titre campagne
    t = doc.add_paragraph(campaign.get("title", "Campagne"))
    _para_spacing(t, before=16, after=6)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.bold = True; t.runs[0].font.size = Pt(16)
    t.runs[0].font.color.rgb = BLUE_PRIMARY; t.runs[0].font.name = "Calibri"

    # Diagnostic
    diag_para = doc.add_paragraph()
    _para_spacing(diag_para, before=4, after=16)
    diag_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = diag_para.add_run(f"● {DIAGNOSTIC_LABEL.get(diagnostic, 'ATTENTION')}")
    dr.font.bold = True; dr.font.size = Pt(11)
    dr.font.color.rgb = DIAGNOSTIC_COLOR.get(diagnostic, ORANGE)
    dr.font.name = "Calibri"

    # Tableau de synthèse couverture
    tbl = doc.add_table(rows=4, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    rows_data = [
        ("Statut", CAMPAIGN_STATUS_FR.get(campaign.get("status", ""), campaign.get("status", ""))),
        ("Période", f'{campaign.get("start_date") or "N/A"}  →  {campaign.get("end_date") or "N/A"}'),
        ("Avancement", f'{stats.get("progress_pct", 0)}%  ({stats.get("done", 0)} / {stats.get("total", 0)} procédures)'),
        ("Généré le", now_str),
    ]
    for i, (label, value) in enumerate(rows_data):
        row = tbl.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        _set_cell_bg(row.cells[0], BLUE_PALE)
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.name = "Calibri"
            _set_cell_border(cell)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = BLUE_PRIMARY

    # Description
    if campaign.get("description"):
        desc = doc.add_paragraph(campaign["description"])
        _para_spacing(desc, before=12, after=4)
        desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc.runs[0].font.size = Pt(10); desc.runs[0].font.italic = True
        desc.runs[0].font.color.rgb = GREY_TEXT; desc.runs[0].font.name = "Calibri"

    doc.add_page_break()

    # ── SOMMAIRE ─────────────────────────────────────────────────
    s = doc.add_paragraph("SOMMAIRE")
    _para_spacing(s, before=0, after=8)
    s.runs[0].font.bold = True; s.runs[0].font.size = Pt(13)
    s.runs[0].font.color.rgb = BLUE_PRIMARY; s.runs[0].font.name = "Calibri"
    _add_toc(doc)
    doc.add_page_break()

    # ── 1. RÉSUMÉ EXÉCUTIF ────────────────────────────────────────
    _heading(doc, "1. Résumé exécutif", level=1)
    _body(doc, narrative.get("resume_executif", ""))

    # ── 2. TABLEAU DE BORD ────────────────────────────────────────
    _heading(doc, "2. Tableau de bord", level=1)

    tbl2 = doc.add_table(rows=2, cols=4)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2.style = "Table Grid"
    headers = ["Total", "Formalisées", "En cours", "En attente"]
    values = [
        str(stats.get("total", 0)),
        str(stats.get("done", 0)),
        str(stats.get("in_progress", 0)),
        str(stats.get("pending", 0)),
    ]
    for i, h in enumerate(headers):
        cell = tbl2.rows[0].cells[i]
        cell.text = h
        _set_cell_bg(cell, BLUE_PRIMARY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.bold = True; p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = WHITE; p.runs[0].font.name = "Calibri"
        _set_cell_border(cell)
    for i, v in enumerate(values):
        cell = tbl2.rows[1].cells[i]
        cell.text = v
        _set_cell_bg(cell, BLUE_PALE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.bold = True; p.runs[0].font.size = Pt(14)
        p.runs[0].font.color.rgb = BLUE_PRIMARY; p.runs[0].font.name = "Calibri"
        _set_cell_border(cell)

    # Barre de progression textuelle
    pct = stats.get("progress_pct", 0)
    filled = round(pct / 5)
    bar = "█" * filled + "░" * (20 - filled)
    prog = doc.add_paragraph(f"\n  {bar}  {pct}%")
    _para_spacing(prog, before=6, after=10)
    prog.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prog.runs[0].font.size = Pt(12); prog.runs[0].font.name = "Courier New"
    prog.runs[0].font.color.rgb = BLUE_LIGHT

    # ── 3. ANALYSE DE L'AVANCEMENT ────────────────────────────────
    _heading(doc, "3. Analyse de l'avancement", level=1)
    _body(doc, narrative.get("analyse_avancement", ""))

    # ── 4. POINTS DE BLOCAGE ──────────────────────────────────────
    _heading(doc, "4. Points de blocage", level=1)
    blocages = narrative.get("points_blocage") or []
    if blocages:
        for i, b in enumerate(blocages, 1):
            sub = doc.add_paragraph(f"{i}. {b.get('titre', '')}")
            _para_spacing(sub, before=6, after=2)
            sub.runs[0].font.bold = True; sub.runs[0].font.size = Pt(10.5)
            sub.runs[0].font.color.rgb = BLUE_PRIMARY; sub.runs[0].font.name = "Calibri"
            desc_p = _body(doc, b.get("description", ""))
            pids = b.get("procedures_impactees") or []
            if pids:
                imp = _body(doc, "Procédures impactées : " + ", ".join(pids))
                imp.runs[0].font.italic = True
    else:
        _body(doc, "Aucun point de blocage majeur identifié.")

    # ── 5. DÉTAIL DES PROCÉDURES ──────────────────────────────────
    _heading(doc, "5. Détail des procédures", level=1)
    if procs:
        col_widths = [Inches(0.4), Inches(1.2), Inches(2.4), Inches(1.1), Inches(1.0), Inches(0.9)]
        tbl3 = doc.add_table(rows=1, cols=6)
        tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl3.style = "Table Grid"
        # Largeurs
        for j, w in enumerate(col_widths):
            for cell in [row.cells[j] for row in tbl3.rows]:
                cell.width = w

        headers3 = ["#", "Référence", "Nom", "Statut", "Responsable", "Étapes"]
        hdr_row = tbl3.rows[0]
        for j, h in enumerate(headers3):
            c = hdr_row.cells[j]
            c.text = h
            _set_cell_bg(c, BLUE_PRIMARY)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if c.paragraphs[0].runs:
                r = c.paragraphs[0].runs[0]
                r.font.bold = True; r.font.size = Pt(9)
                r.font.color.rgb = WHITE; r.font.name = "Calibri"
            _set_cell_border(c)

        for idx, p in enumerate(procs, 1):
            row = tbl3.add_row()
            lc = p.get("lifecycle") or {}
            stages_str = f'{lc.get("stages_done", 0)}/{lc.get("stages_total", 6)}'
            row_bg = GREY_LIGHT if idx % 2 == 0 else WHITE
            data = [
                str(idx),
                p.get("procedure_ref") or "—",
                p.get("procedure_nom") or "—",
                PROC_STATUS_FR.get(p.get("status", "pending"), p.get("status", "")),
                p.get("responsible_name") or "—",
                stages_str,
            ]
            for j, val in enumerate(data):
                c = row.cells[j]
                c.text = val
                _set_cell_bg(c, row_bg)
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 2 else WD_ALIGN_PARAGRAPH.LEFT
                if c.paragraphs[0].runs:
                    r = c.paragraphs[0].runs[0]
                    r.font.size = Pt(9); r.font.name = "Calibri"
                    r.font.color.rgb = GREY_TEXT
                _set_cell_border(c)
    else:
        _body(doc, "Aucune procédure dans ce projet.")

    # ── 6. RECOMMANDATIONS ────────────────────────────────────────
    _heading(doc, "6. Recommandations", level=1)
    recs = narrative.get("recommandations") or []
    PRIO_COLORS = {"haute": RED, "moyenne": ORANGE, "faible": GREEN}
    if recs:
        rec_tbl = doc.add_table(rows=1, cols=4)
        rec_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        rec_tbl.style = "Table Grid"
        for j, h in enumerate(["Priorité", "Action", "Description", "Responsable"]):
            c = rec_tbl.rows[0].cells[j]
            c.text = h
            _set_cell_bg(c, BLUE_PRIMARY)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if c.paragraphs[0].runs:
                r2 = c.paragraphs[0].runs[0]
                r2.font.bold = True; r2.font.size = Pt(9)
                r2.font.color.rgb = WHITE; r2.font.name = "Calibri"
            _set_cell_border(c)
        for rec in recs:
            row = rec_tbl.add_row()
            prio = rec.get("priorite", "moyenne")
            data = [prio.upper(), rec.get("action", ""), rec.get("description", ""), rec.get("responsable", "")]
            for j, val in enumerate(data):
                c = row.cells[j]
                c.text = val
                _set_cell_bg(c, GREY_LIGHT if j % 2 else WHITE)
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
                if c.paragraphs[0].runs:
                    r2 = c.paragraphs[0].runs[0]
                    r2.font.size = Pt(9); r2.font.name = "Calibri"
                    if j == 0:
                        r2.font.bold = True
                        r2.font.color.rgb = PRIO_COLORS.get(prio, ORANGE)
                    else:
                        r2.font.color.rgb = GREY_TEXT
                _set_cell_border(c)
    else:
        _body(doc, "Aucune recommandation spécifique émise.")

    # ── 7. CONCLUSION ──────────────────────────────────────────────
    _heading(doc, "7. Conclusion", level=1)
    _body(doc, narrative.get("conclusion", ""))

    # ── PIED DE PAGE ───────────────────────────────────────────────
    footer_para = doc.add_paragraph()
    _para_spacing(footer_para, before=24, after=0)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top"); top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4"); top.set(qn("w:color"), _hex(BLUE_LIGHT))
    hr.append(top)
    footer_para._p.get_or_add_pPr().append(hr)
    fr = footer_para.add_run(f"Document généré automatiquement par ProcessMate · {now_str}")
    fr.font.size = Pt(8); fr.font.italic = True
    fr.font.color.rgb = GREY_TEXT; fr.font.name = "Calibri"

    # Sérialise en mémoire
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── Endpoint ─────────────────────────────────────────────────

@router.get("/{campaign_id}/report")
async def download_campaign_report(campaign_id: str):
    """Génère et retourne le rapport Word (.docx) de la campagne."""
    try:
        campaign = _get_campaign_full(campaign_id)
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Erreur récupération campagne %s: %s", campaign_id, e)
        raise HTTPException(500, "Erreur lors de la récupération de la campagne")

    try:
        narrative = await generate_campaign_narrative(campaign)
    except Exception as e:
        logger.error("Erreur Gemini rapport %s: %s", campaign_id, e)
        raise HTTPException(500, f"Erreur lors de la génération IA : {e}")

    try:
        docx_bytes = _build_docx(campaign, narrative)
    except Exception as e:
        logger.error("Erreur génération Word %s: %s", campaign_id, e)
        raise HTTPException(500, "Erreur lors de la construction du document Word")

    title_slug = (campaign.get("title") or "rapport").replace(" ", "_")[:40]
    filename = f"rapport_projet_{title_slug}.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─── Rapport Portfolio (tous les projets) ─────────────────────

@router.get("/portfolio/report")
async def download_portfolio_report():
    """Génère un rapport Word consolidé de TOUS les projets de formalisation."""
    db = get_supabase()

    # Récupère tous les projets avec leurs stats
    res = db.table("formalization_campaigns").select("*").order("created_at", desc=True).execute()
    projects = res.data or []
    if not projects:
        raise HTTPException(404, "Aucun projet trouvé")

    # Enrichit chaque projet avec ses stats
    for p in projects:
        cp = db.table("campaign_procedures").select("status, procedure_id").eq("campaign_id", p["id"]).execute()
        procs = cp.data or []
        total = len(procs)
        done = sum(1 for x in procs if x.get("status") in ("formalized", "validated"))
        in_progress = sum(1 for x in procs if x.get("status") == "in_progress")
        p["stats"] = {
            "total": total, "done": done,
            "in_progress": in_progress,
            "pending": total - done - in_progress,
            "progress_pct": round((done / total * 100) if total else 0),
        }
        p["procedures"] = procs

    try:
        narrative = await _generate_portfolio_narrative(projects)
    except Exception as e:
        logger.error("Erreur Gemini portfolio: %s", e)
        raise HTTPException(500, f"Erreur IA : {e}")

    try:
        docx_bytes = _build_portfolio_docx(projects, narrative)
    except Exception as e:
        logger.error("Erreur Word portfolio: %s", e)
        raise HTTPException(500, "Erreur lors de la construction du rapport")

    now = datetime.utcnow().strftime("%Y%m%d_%H%M")
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="portfolio_projets_{now}.docx"'},
    )


# ─── Narrative portfolio (Gemini) ─────────────────────────────

async def _generate_portfolio_narrative(projects: list) -> dict:
    import asyncio, json, re, os
    from manager.model_manager import GeminiModelManager

    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        raise ValueError("GOOGLE_API_KEY non configurée")

    total_projs   = len(projects)
    active        = [p for p in projects if p.get("status") == "active"]
    blocked       = [p for p in projects if p.get("status") == "blocked"]
    on_hold       = [p for p in projects if p.get("status") == "on_hold"]
    completed     = [p for p in projects if p.get("status") == "completed"]
    total_procs   = sum(p["stats"]["total"] for p in projects)
    done_procs    = sum(p["stats"]["done"]  for p in projects)
    avg_pct       = round(sum(p["stats"]["progress_pct"] for p in projects) / total_projs) if total_projs else 0

    proj_lines = "\n".join(
        f"  - {p.get('title','?')} | statut: {p.get('status','?')} | "
        f"{p['stats']['total']} proc. | {p['stats']['progress_pct']}% avancement"
        for p in projects
    )

    prompt = f"""Tu es un expert en pilotage de projets de formalisation de procédures bancaires.
Tu rédiges un rapport de synthèse du portfolio de projets de formalisation.

DONNÉES DU PORTFOLIO :
- Total projets : {total_projs}
- Actifs : {len(active)} | Bloqués : {len(blocked)} | En pause : {len(on_hold)} | Terminés : {len(completed)}
- Total procédures : {total_procs} | Formalisées : {done_procs} | Avancement moyen : {avg_pct}%

DÉTAIL DES PROJETS :
{proj_lines}

Retourne uniquement du JSON valide (sans markdown) :
{{
  "diagnostic_global": "positif" | "acceptable" | "attention" | "critique",
  "synthese": "Paragraphe de synthèse du portfolio (3-5 phrases).",
  "analyse": "Analyse de l'avancement global, tendances, points forts et retards (3-4 phrases).",
  "points_critiques": ["Point critique 1", "Point critique 2"],
  "recommandations": [
    {{"priorite": "haute"|"moyenne"|"faible", "action": "Action courte", "description": "Détail actionnable"}}
  ],
  "conclusion": "Conclusion et prochaines étapes (2-3 phrases)."
}}"""

    manager = GeminiModelManager(api_key)

    async def _task(model_name: str):
        model = manager.get_model(model_name)
        return await asyncio.wait_for(
            asyncio.to_thread(model.generate_content, model=model_name, contents=[prompt]),
            timeout=120,
        )

    result = await manager.execute_with_fallback(_task, "Rapport portfolio")
    if not result.get("success"):
        raise ValueError(result.get("message") or "Erreur Gemini")

    raw = getattr(result["result"], "text", "") or ""
    raw = re.sub(r"^```(?:json)?\s*", "", raw.strip())
    raw = re.sub(r"\s*```$", "", raw)
    try:
        return json.loads(raw)
    except Exception:
        return {"diagnostic_global": "attention", "synthese": raw[:500], "analyse": "",
                "points_critiques": [], "recommandations": [], "conclusion": ""}


# ─── Word portfolio ────────────────────────────────────────────

def _build_portfolio_docx(projects: list, narrative: dict) -> bytes:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3);  sec.right_margin = Cm(2.5)

    _enable_auto_update_fields(doc)

    now_str  = datetime.utcnow().strftime("%d/%m/%Y à %H:%M")
    diag     = narrative.get("diagnostic_global", "attention")
    diag_cfg = {"positif": (GREEN,"POSITIF"), "acceptable": (BLUE_LIGHT,"ACCEPTABLE"),
                "attention": (ORANGE,"ATTENTION REQUISE"), "critique": (RED,"CRITIQUE")}
    diag_color, diag_label = diag_cfg.get(diag, (ORANGE, "ATTENTION REQUISE"))

    total_projs = len(projects)
    total_procs = sum(p["stats"]["total"] for p in projects)
    done_procs  = sum(p["stats"]["done"]  for p in projects)
    avg_pct     = round(sum(p["stats"]["progress_pct"] for p in projects) / total_projs) if total_projs else 0

    # ── PAGE DE GARDE ────────────────────────────────────────────
    cover = doc.add_paragraph()
    _para_spacing(cover, before=40, after=8)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = cover._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), _hex(BLUE_PRIMARY)); pPr.append(shd)
    _r = cover.add_run("  RAPPORT PORTFOLIO — PROJETS DE FORMALISATION  ")
    _r.font.bold = True; _r.font.size = Pt(16); _r.font.color.rgb = WHITE; _r.font.name = "Calibri"

    diag_p = doc.add_paragraph()
    _para_spacing(diag_p, before=10, after=16)
    diag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = diag_p.add_run(f"● {diag_label}")
    dr.font.bold = True; dr.font.size = Pt(11); dr.font.color.rgb = diag_color; dr.font.name = "Calibri"

    # KPIs page de garde
    tbl = doc.add_table(rows=2, cols=4); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER; tbl.style = "Table Grid"
    for j, h in enumerate(["Projets", "Procédures", "Formalisées", "Avancement moyen"]):
        c = tbl.rows[0].cells[j]; c.text = h
        _set_cell_bg(c, BLUE_PRIMARY); _set_cell_border(c)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].font.bold = True
            c.paragraphs[0].runs[0].font.color.rgb = WHITE
            c.paragraphs[0].runs[0].font.size = Pt(9)
    for j, v in enumerate([str(total_projs), str(total_procs), str(done_procs), f"{avg_pct}%"]):
        c = tbl.rows[1].cells[j]; c.text = v
        _set_cell_bg(c, BLUE_PALE); _set_cell_border(c)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].font.bold = True
            c.paragraphs[0].runs[0].font.color.rgb = BLUE_PRIMARY
            c.paragraphs[0].runs[0].font.size = Pt(14)

    gd = doc.add_paragraph(f"\nGénéré le {now_str}")
    _para_spacing(gd, before=8); gd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if gd.runs: gd.runs[0].font.size = Pt(9); gd.runs[0].font.color.rgb = GREY_TEXT
    doc.add_page_break()

    # ── SOMMAIRE ─────────────────────────────────────────────────
    s = doc.add_paragraph("SOMMAIRE")
    _para_spacing(s, before=0, after=8)
    s.runs[0].font.bold = True; s.runs[0].font.size = Pt(13)
    s.runs[0].font.color.rgb = BLUE_PRIMARY; s.runs[0].font.name = "Calibri"
    _add_toc(doc)
    doc.add_page_break()

    # ── 1. SYNTHÈSE ───────────────────────────────────────────────
    _heading(doc, "1. Synthèse du portfolio", level=1)
    _body(doc, narrative.get("synthese", ""))

    # ── 2. ANALYSE ────────────────────────────────────────────────
    _heading(doc, "2. Analyse de l'avancement", level=1)
    _body(doc, narrative.get("analyse", ""))

    # ── 3. VUE CONSOLIDÉE DES PROJETS ─────────────────────────────
    _heading(doc, "3. Vue consolidée des projets", level=1)
    status_order = {"blocked": 0, "on_hold": 1, "active": 2, "draft": 3, "completed": 4, "archived": 5}
    sorted_projs = sorted(projects, key=lambda p: status_order.get(p.get("status",""), 9))

    tbl3 = doc.add_table(rows=1, cols=5); tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER; tbl3.style = "Table Grid"
    for j, h in enumerate(["Projet", "Statut", "Procédures", "Avancement", "Échéance"]):
        c = tbl3.rows[0].cells[j]; c.text = h
        _set_cell_bg(c, BLUE_PRIMARY); _set_cell_border(c)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].font.bold = True
            c.paragraphs[0].runs[0].font.color.rgb = WHITE
            c.paragraphs[0].runs[0].font.size = Pt(9)

    for idx, p in enumerate(sorted_projs):
        row = tbl3.add_row()
        bg = GREY_LIGHT if idx % 2 == 0 else WHITE
        status_label = PROJECT_STATUS_FR.get(p.get("status", ""), p.get("status", ""))
        pct = p["stats"]["progress_pct"]
        bar = "█" * round(pct / 10) + "░" * (10 - round(pct / 10))
        vals = [
            p.get("title", "—"),
            status_label,
            f'{p["stats"]["total"]} proc. · {p["stats"]["done"]} faites',
            f'{bar} {pct}%',
            p.get("end_date", "—")[:10] if p.get("end_date") else "—",
        ]
        for j, val in enumerate(vals):
            c = row.cells[j]; c.text = val
            _set_cell_bg(c, bg); _set_cell_border(c)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            if c.paragraphs[0].runs:
                r2 = c.paragraphs[0].runs[0]; r2.font.size = Pt(8); r2.font.name = "Calibri"
                if j == 3: r2.font.name = "Courier New"; r2.font.size = Pt(8)
                # Couleur statut
                if j == 1:
                    st = p.get("status", "")
                    r2.font.color.rgb = RED if st == "blocked" else (ORANGE if st == "on_hold" else GREY_TEXT)
                    r2.font.bold = st in ("blocked", "on_hold")
                else:
                    r2.font.color.rgb = GREY_TEXT

    # ── 4. POINTS CRITIQUES ───────────────────────────────────────
    _heading(doc, "4. Points critiques", level=1)
    critiques = narrative.get("points_critiques") or []
    if critiques:
        for pt in critiques:
            p_c = doc.add_paragraph(f"• {pt}")
            _para_spacing(p_c, before=2, after=2)
            if p_c.runs: p_c.runs[0].font.size = Pt(10); p_c.runs[0].font.color.rgb = GREY_TEXT
    else:
        _body(doc, "Aucun point critique identifié.")

    # ── 5. RECOMMANDATIONS ────────────────────────────────────────
    _heading(doc, "5. Recommandations", level=1)
    recs = narrative.get("recommandations") or []
    PRIO_COLORS = {"haute": RED, "moyenne": ORANGE, "faible": GREEN}
    if recs:
        rec_tbl = doc.add_table(rows=1, cols=3); rec_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER; rec_tbl.style = "Table Grid"
        for j, h in enumerate(["Priorité", "Action", "Description"]):
            c = rec_tbl.rows[0].cells[j]; c.text = h
            _set_cell_bg(c, BLUE_PRIMARY); _set_cell_border(c)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.bold = True
                c.paragraphs[0].runs[0].font.color.rgb = WHITE
                c.paragraphs[0].runs[0].font.size = Pt(9)
        for rec in recs:
            row = rec_tbl.add_row()
            prio = rec.get("priorite", "moyenne")
            for j, val in enumerate([prio.upper(), rec.get("action", ""), rec.get("description", "")]):
                c = row.cells[j]; c.text = val; _set_cell_border(c)
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
                if c.paragraphs[0].runs:
                    r2 = c.paragraphs[0].runs[0]; r2.font.size = Pt(9); r2.font.name = "Calibri"
                    if j == 0: r2.font.bold = True; r2.font.color.rgb = PRIO_COLORS.get(prio, ORANGE)
                    else: r2.font.color.rgb = GREY_TEXT
    else:
        _body(doc, "Aucune recommandation émise.")

    # ── 6. CONCLUSION ──────────────────────────────────────────────
    _heading(doc, "6. Conclusion", level=1)
    _body(doc, narrative.get("conclusion", ""))

    # Pied de page
    footer = doc.add_paragraph()
    _para_spacing(footer, before=24); footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr2 = OxmlElement("w:pBdr"); tp2 = OxmlElement("w:top")
    tp2.set(qn("w:val"), "single"); tp2.set(qn("w:sz"), "4"); tp2.set(qn("w:color"), _hex(BLUE_LIGHT))
    hr2.append(tp2); footer._p.get_or_add_pPr().append(hr2)
    fr = footer.add_run(f"Rapport portfolio généré par ProcessMate · {now_str}")
    fr.font.size = Pt(8); fr.font.italic = True; fr.font.color.rgb = GREY_TEXT; fr.font.name = "Calibri"

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()
