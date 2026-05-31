"""Generate ProcessMate requirements Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "ProcessMate_Exigences_v2.docx")
SVG_DIR = os.path.dirname(__file__)

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color="1E293B"):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return h

def add_paragraph(doc, text, bold=False, color=None, size=11, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_bullet(doc, text, bold_prefix=None, indent=1):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(indent)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)

def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

def info_box(doc, label, text, bg="FEF3C7", label_color="92400E", text_color="78350F"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    if label:
        r1 = p.add_run(label + "  ")
        r1.bold = True
        r1.font.color.rgb = RGBColor.from_string(label_color)
        r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor.from_string(text_color)
    r2.font.size = Pt(11)
    doc.add_paragraph()

def section_box(doc, title, items, title_bg="3B82F6"):
    """items = list of (label, value) tuples"""
    table = doc.add_table(rows=1 + len(items), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header
    hdr = table.cell(0, 0).merge(table.cell(0, 1))
    set_cell_bg(hdr, title_bg)
    p = hdr.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # rows
    for i, (lbl, val) in enumerate(items, 1):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        set_cell_bg(c0, "F8FAFC")
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(lbl)
        r0.bold = True
        r0.font.size = Pt(10)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(10)
    doc.add_paragraph()

# ── document ─────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── COVER ────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PROCESSMATE")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Exigences Techniques — Installation & Déploiement")
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()
info_box(doc,
    "Objectif de ce document :",
    "Exposer les conditions nécessaires pour faire fonctionner ProcessMate, "
    "que ce soit en local sur une machine ou en production via un déploiement cloud. "
    "Ce document permet de comparer les deux approches et d'éclairer le choix d'infrastructure.",
    bg="DBEAFE", label_color="1E40AF", text_color="1E3A5F")

doc.add_page_break()

# ── SECTION 1 : VUE D'ENSEMBLE ───────────────────────────────────────────────
add_heading(doc, "1. Vue d'ensemble du projet")

add_paragraph(doc,
    "ProcessMate est une plateforme web de modélisation de processus métier (BPMN). "
    "Elle est composée d'un frontend Next.js et de trois APIs backend indépendantes "
    "développées en Python/FastAPI, chacune exposée sur un port distinct.",
    size=11)

doc.add_paragraph()

section_box(doc, "Composants de l'application", [
    ("Frontend",        "Next.js 16 + React 19 — Interface utilisateur (port 3000)"),
    ("Clinic API",      "FastAPI — BPMN, OCR, Chat IA, analyse (port 8002)"),
    ("SCV Maker API",   "FastAPI — Génération et validation de SCV (port 8001)"),
    ("SFD Generator",   "FastAPI — Documents de spécification (port 8004)"),
    ("Base de données", "Supabase PostgreSQL (cloud) — sessions, fichiers, workflows"),
    ("IA / LLM",        "Google Gemini (via GOOGLE_API_KEY) — moteur d'analyse IA"),
], title_bg="1E293B")

doc.add_page_break()

# ── SECTION 2 : INSTALLATION LOCALE ─────────────────────────────────────────
add_heading(doc, "2. Installation Locale")

p = doc.add_paragraph()
r = p.add_run("Durée estimée : 2 à 4 heures par machine")
r.bold = True
r.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
r.font.size = Pt(12)

info_box(doc,
    "Important :",
    "Chaque utilisateur souhaitant lancer l'application doit répéter l'intégralité "
    "de ces étapes sur sa propre machine. Il n'existe pas d'installeur automatique.",
    bg="FEE2E2", label_color="991B1B", text_color="7F1D1D")

# 2.1 Prérequis système
add_heading(doc, "2.1 Prérequis système", level=2)
section_box(doc, "Environnement requis", [
    ("Python",        "Version 3.10 ou supérieure"),
    ("Node.js",       "Version 18 LTS ou supérieure"),
    ("npm",           "Version 9+ (inclus avec Node.js)"),
    ("RAM",           "8 Go minimum recommandés"),
    ("Connexion",     "Internet requise (APIs Supabase et Gemini sont cloud)"),
    ("Terminaux",     "4 fenêtres de terminal ouvertes simultanément"),
], title_bg="3B82F6")

# 2.2 Dépendances backend
add_heading(doc, "2.2 Dépendances Backend Python", level=2)
add_paragraph(doc, "Commande d'installation :", bold=True)
add_code_block(doc, ["pip install -r requirements.txt"])

section_box(doc, "Packages Python principaux (requirements.txt)", [
    ("FastAPI 0.115",        "Framework web — toutes les APIs"),
    ("Uvicorn 0.30",         "Serveur ASGI pour FastAPI"),
    ("Pandas 2.2",           "Manipulation de données Excel/CSV"),
    ("OpenPyXL 3.1",         "Lecture/écriture fichiers .xlsx"),
    ("Pillow 10.4",          "Traitement d'images"),
    ("Matplotlib 3.9",       "Génération de graphiques"),
    ("Pydantic 2.9",         "Validation des données"),
    ("python-multipart 0.0.6","Upload de fichiers"),
    ("Supabase-py",          "Client base de données Supabase"),
    ("google-generativeai",  "API Google Gemini (IA)"),
    ("aiohttp / httpx",      "Requêtes HTTP asynchrones"),
    ("networkx 3.2",         "Graphes pour BPMN"),
], title_bg="7C3AED")

# 2.3 Dépendances frontend
add_heading(doc, "2.3 Dépendances Frontend Node.js", level=2)
add_paragraph(doc, "Commande d'installation (dans le dossier frontend/) :", bold=True)
add_code_block(doc, ["cd frontend", "npm install"])

info_box(doc,
    "Volume :",
    "454 packages npm installés. Le dossier node_modules/ peut atteindre 800 Mo à 1,2 Go.",
    bg="FEF3C7", label_color="92400E", text_color="78350F")

section_box(doc, "Packages npm principaux (package.json)", [
    ("Next.js 16 + React 19",     "Framework frontend principal"),
    ("@mui/material 7.3",         "Composants UI Material Design"),
    ("Tailwind CSS 4",            "Styles utilitaires CSS"),
    ("bpmn-js 18.8",              "Visualiseur et éditeur BPMN"),
    ("ReactFlow 11.11",           "Diagrammes interactifs"),
    ("Mermaid 11.12",             "Diagrammes textuels"),
    ("Handsontable 16",           "Grille de données type Excel"),
    ("XLSX 0.18",                 "Export/import fichiers Excel"),
    ("@supabase/supabase-js 2.75","Client Supabase (auth + BDD)"),
    ("Lucide React 0.545",        "Icônes"),
    ("Zustand 5",                 "Gestion d'état global"),
], title_bg="059669")

# 2.4 Clés API
add_heading(doc, "2.4 Clés API requises", level=2)

info_box(doc,
    "Attention :",
    "Ces clés ne peuvent pas être partagées librement. "
    "Elles sont liées à des comptes cloud (Google Cloud, Supabase) et doivent être "
    "obtenues séparément par l'administrateur.",
    bg="FEE2E2", label_color="991B1B", text_color="7F1D1D")

section_box(doc, "Clés API et credentials nécessaires", [
    ("GOOGLE_API_KEY",              "Clé API Google Gemini — OBLIGATOIRE"),
    ("SUPABASE_URL (backend)",      "URL du projet Supabase côté serveur"),
    ("SUPABASE_SERVICE_KEY",        "Clé de service Supabase (accès admin DB)"),
    ("NEXT_PUBLIC_SUPABASE_URL",    "URL Supabase côté frontend"),
    ("NEXT_PUBLIC_SUPABASE_ANON_KEY","Clé publique Supabase (authentification)"),
], title_bg="DC2626")

# 2.5 Fichiers .env
add_heading(doc, "2.5 Fichiers de configuration .env", level=2)
add_paragraph(doc,
    "Trois fichiers d'environnement doivent être créés manuellement "
    "(ils ne sont pas inclus dans le code source pour des raisons de sécurité) :")

add_bullet(doc, " clinic/.env", bold_prefix="Fichier 1 :")
add_code_block(doc, [
    "GOOGLE_API_KEY=AIzaSy...",
    "SUPABASE_URL=https://agotdktomrhlopryqiow.supabase.co",
    "SUPABASE_SERVICE_KEY=eyJhbGc...",
    "FRONTEND_URL=http://localhost:3000",
    "IS_PRODUCTION=false",
])

add_bullet(doc, " sfd/.env", bold_prefix="Fichier 2 :")
add_code_block(doc, [
    "GOOGLE_API_KEY=AIzaSy...",
    "PORT=8004",
    "FRONTEND_URL=http://localhost:3000",
])

add_bullet(doc, " frontend/.env.local", bold_prefix="Fichier 3 :")
add_code_block(doc, [
    "NEXT_PUBLIC_API_URL=http://localhost:8002",
    "NEXT_PUBLIC_SCV_API_URL=http://localhost:8001",
    "NEXT_PUBLIC_SFD_API_URL=http://localhost:8004",
    "NEXT_PUBLIC_SUPABASE_URL=https://owafxzgjsvrdcmbscgmg.supabase.co",
    "NEXT_PUBLIC_SUPABASE_ANON_KEY=eyJhbGc...",
])

doc.add_paragraph()

# 2.6 Démarrage
add_heading(doc, "2.6 Procédure de démarrage", level=2)
add_paragraph(doc,
    "À chaque session de travail, 4 terminaux doivent être ouverts simultanément :")

section_box(doc, "Commandes de démarrage (dans l'ordre)", [
    ("Terminal 1 — Frontend",   "cd frontend && npm run dev          → http://localhost:3000"),
    ("Terminal 2 — Clinic API", "python clinic/main.py               → http://localhost:8002"),
    ("Terminal 3 — SCV API",    "python backend/main.py              → http://localhost:8001"),
    ("Terminal 4 — SFD API",    "python sfd/main.py                  → http://localhost:8004"),
], title_bg="F59E0B")

doc.add_page_break()

# ── SECTION 3 : DÉPLOIEMENT ──────────────────────────────────────────────────
add_heading(doc, "3. Exigences de Déploiement")

p = doc.add_paragraph()
r = p.add_run("Durée estimée : 1 journée de travail")
r.bold = True
r.font.color.rgb = RGBColor(0xEA, 0x58, 0x0C)
r.font.size = Pt(12)

add_paragraph(doc,
    "Le déploiement transforme ProcessMate en service accessible depuis n'importe quel "
    "navigateur, sans installation sur les postes utilisateurs. "
    "Il nécessite la mise en place de quatre points clés détaillés ci-dessous.", size=11)
doc.add_paragraph()

# 3.1 Ports
add_heading(doc, "3.1 Gestion des 4 ports backend", level=2)
add_paragraph(doc,
    "Chacun des trois backends FastAPI et le frontend Next.js s'exécutent sur "
    "un port distinct. En production, chaque service doit être déployé et exposé "
    "séparément (ex : 4 applications Render.com, ou 4 processus derrière un reverse proxy).")

section_box(doc, "Ports et services", [
    ("Port 3000", "Frontend Next.js           — npm start après build"),
    ("Port 8002", "Clinic API (FastAPI)        — service principal BPMN + IA"),
    ("Port 8001", "SCV Maker API (FastAPI)     — génération de SCV"),
    ("Port 8004", "SFD Generator API (FastAPI) — documents de spécification"),
], title_bg="0F172A")

add_paragraph(doc,
    "En production sur Render.com, chaque service est une application distincte "
    "avec son propre URL (ex : processmate-back.onrender.com). "
    "Un reverse proxy (Nginx ou équivalent) peut être utilisé pour unifier sous un seul domaine.",
    color="475569")

# 3.2 CORS
add_heading(doc, "3.2 Gestion du CORS", level=2)
add_paragraph(doc,
    "Le CORS (Cross-Origin Resource Sharing) doit être explicitement configuré "
    "sur chaque backend pour autoriser les requêtes provenant du frontend déployé.")

info_box(doc,
    "Principe :",
    "Si FRONTEND_URL n'est pas correctement défini côté backend, "
    "le navigateur bloquera toutes les requêtes API — l'application sera inutilisable.",
    bg="FEF3C7", label_color="92400E", text_color="78350F")

section_box(doc, "Variables CORS à définir sur chaque backend", [
    ("FRONTEND_URL",   "URL complète du frontend déployé (ex : https://processmate.vercel.app)"),
    ("IS_PRODUCTION",  "Mettre à true pour activer le CORS strict en production"),
    ("Origines locales","localhost:3000/3001/3002 sont inclus automatiquement en mode dev"),
    ("allow_methods",  '["*"] — toutes méthodes HTTP autorisées'),
    ("allow_headers",  '["*"] — tous headers autorisés (dont Authorization)'),
], title_bg="CA8A04")

add_paragraph(doc,
    "Le code CORS est déjà implémenté dans chaque backend. "
    "Seule la variable d'environnement FRONTEND_URL doit être renseignée "
    "après déploiement du frontend.", color="475569")

# 3.3 Build et chemins
add_heading(doc, "3.3 Build frontend et gestion des chemins", level=2)
add_paragraph(doc,
    "En production, le frontend Next.js doit être compilé avant déploiement. "
    "Cette étape transforme le code TypeScript/React en fichiers optimisés.")

section_box(doc, "Étapes de build frontend", [
    ("1. npm run build",    "Compile et optimise le code — génère le dossier .next/"),
    ("2. npm start",        "Lance le serveur de production (port 3000 par défaut)"),
    ("3. Variables .env",   "Le fichier .env.production doit être présent avant le build"),
    ("4. URLs figées",      "Les NEXT_PUBLIC_* sont intégrées au build — tout changement nécessite un rebuild"),
], title_bg="6D28D9")

info_box(doc,
    "Point d'attention :",
    "Les variables d'environnement NEXT_PUBLIC_* sont incorporées à la compilation. "
    "Si les URLs des backends changent après le build, il faut recompiler le frontend.",
    bg="EDE9FE", label_color="4C1D95", text_color="5B21B6")

# 3.4 Liaison frontend-backend
add_heading(doc, "3.4 Liaison Frontend ↔ Backend", level=2)
add_paragraph(doc,
    "Le fichier frontend/src/lib/api-config.ts est le point central qui "
    "définit vers quels serveurs le frontend envoie ses requêtes. "
    "Il lit les variables d'environnement NEXT_PUBLIC_* pour construire les URLs.")

add_paragraph(doc, "Exemple de configuration production :", bold=True)
add_code_block(doc, [
    "// frontend/.env.production",
    "NEXT_PUBLIC_API_URL=https://processmate-back.onrender.com",
    "NEXT_PUBLIC_SCV_API_URL=https://scv-maker-api.onrender.com",
    "NEXT_PUBLIC_SFD_API_URL=https://sfd-generator-api.onrender.com",
    "NEXT_PUBLIC_SUPABASE_URL=https://owafxzgjsvrdcmbscgmg.supabase.co",
    "NEXT_PUBLIC_SUPABASE_ANON_KEY=eyJhbGc...",
])

doc.add_paragraph()

section_box(doc, "Checklist liaison frontend-backend", [
    ("Étape 1", "Déployer les 3 backends — noter leurs URLs publiques"),
    ("Étape 2", "Renseigner ces URLs dans frontend/.env.production"),
    ("Étape 3", "Lancer npm run build avec les bonnes variables"),
    ("Étape 4", "Configurer FRONTEND_URL sur chaque backend = URL du frontend déployé"),
    ("Étape 5", "Tester chaque endpoint API depuis le frontend en production"),
    ("Étape 6", "Vérifier les logs CORS si des requêtes échouent"),
], title_bg="0891B2")

doc.add_page_break()

# ── SECTION 4 : COMPARAISON ──────────────────────────────────────────────────
add_heading(doc, "4. Comparaison : Local vs Déploiement")

compare_data = [
    ("Critère",               "Installation Locale",                  "Déploiement Cloud"),
    ("Temps de mise en place","2–4 heures par machine",               "1 journée (une seule fois)"),
    ("Maintenance",           "Sur chaque poste utilisateur",         "Centralisée"),
    ("Accès utilisateurs",    "1 seule machine",                      "Tout navigateur, partout"),
    ("Clés API",              "À gérer sur chaque poste",             "Centralisées côté serveur"),
    ("Mises à jour",          "À répéter sur chaque machine",         "1 déploiement = tous mis à jour"),
    ("4 terminaux",           "À ouvrir manuellement à chaque fois",  "Gérés automatiquement"),
    ("node_modules (800Mo+)", "Sur chaque poste",                     "Sur le serveur uniquement"),
]
table = doc.add_table(rows=len(compare_data), cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (c0, c1, c2) in enumerate(compare_data):
    row = table.rows[i]
    if i == 0:
        for cell in row.cells:
            set_cell_bg(cell, "374151")
        for cell, txt in zip(row.cells, [c0, c1, c2]):
            r = cell.paragraphs[0].add_run(txt)
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)
    else:
        bg = "F8FAFC" if i % 2 == 0 else "FFFFFF"
        set_cell_bg(row.cells[0], "F1F5F9")
        set_cell_bg(row.cells[1], bg)
        set_cell_bg(row.cells[2], bg)
        row.cells[0].paragraphs[0].add_run(c0).font.size = Pt(10)
        r1 = row.cells[1].paragraphs[0].add_run(c1)
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        r2 = row.cells[2].paragraphs[0].add_run(c2)
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x16, 0x66, 0x34)
doc.add_paragraph()

doc.add_paragraph()

info_box(doc,
    "Recommandation :",
    "Le déploiement cloud est fortement recommandé dès que plusieurs personnes "
    "doivent utiliser ProcessMate. Le coût de la mise en place initiale (1 journée) "
    "est nettement inférieur au coût cumulé des installations individuelles.",
    bg="DCFCE7", label_color="14532D", text_color="166534")

# ── SECTION 5 : SCHÉMAS ──────────────────────────────────────────────────────
add_heading(doc, "5. Schémas visuels")

schemas = [
    ("svg1_installation_locale",        "Schéma 1 — Exigences d'installation locale"),
    ("svg2_architecture_deploiement",   "Schéma 2 — Architecture de déploiement (ports, CORS, services)"),
    ("svg3_liaison_frontend_backend",   "Schéma 3 — Liaison Frontend ↔ Backend (variables & build)"),
]

for name, caption in schemas:
    png_path = os.path.join(SVG_DIR, f"{name}.png")
    if os.path.exists(png_path):
        # caption
        p_cap = doc.add_paragraph()
        r_cap = p_cap.add_run(caption)
        r_cap.bold = True
        r_cap.font.size = Pt(11)
        r_cap.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        # image — 16 cm wide, centred
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(png_path, width=Cm(16))
        doc.add_paragraph()
    else:
        add_bullet(doc, f"{name}.png introuvable", bold_prefix="Manquant : ")

# footer note
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Document généré le 18 mai 2026 — ProcessMate v1.0")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.save(OUTPUT)
print(f"Document saved: {OUTPUT}")
