"""
mockup_word_export.py — Génère un document Word avec les maquettes rebrandées.
Chaque maquette = screenshot HTML capturé par Playwright → image dans Word.
"""

import asyncio
import base64
import io
import os
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from playwright.async_api import async_playwright


async def html_to_screenshot(html: str) -> bytes:
    """Rend un HTML en screenshot PNG via Playwright headless."""
    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=True)
        page = await browser.new_page(viewport={"width": 1440, "height": 900})
        await page.set_content(html, wait_until="networkidle")
        await page.wait_for_timeout(1500)
        screenshot = await page.screenshot(full_page=False, type="png")
        await browser.close()
    return screenshot


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    return RGBColor(r, g, b)


def _zero_spacing(para):
    """Supprime l'espacement before/after d'un paragraphe."""
    pPr = para._p.get_or_add_pPr()
    pPPr = OxmlElement("w:spacing")
    pPPr.set(qn("w:before"), "0")
    pPPr.set(qn("w:after"), "0")
    pPPr.set(qn("w:line"), "240")
    pPPr.set(qn("w:lineRule"), "auto")
    pPr.append(pPPr)


async def generate_mockup_word(
    mockups: list[dict],
    client_name: str,
    source_url: str,
    primary_color: str,
    secondary_color: str,
    logo_b64: str | None = None,
    logo_mime: str = "image/png",
) -> bytes:

    doc = Document()

    # ── Mise en page A4 paysage ───────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width    = Cm(29.7)
    section.page_height   = Cm(21.0)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    content_width_cm = 29.7 - 3.0   # 26.7 cm
    # Hauteur dispo par page après titre + URL (en cm) : ~21 - 1.2 - 1.2 - ~2.5 header = ~16 cm
    img_width_inches  = content_width_cm / 2.54       # 10.51"
    img_height_inches = 16.0 / 2.54                   # 6.30" — plafond hauteur

    # ── Page de titre ─────────────────────────────────────────────────────────
    if logo_b64:
        logo_bytes  = base64.b64decode(logo_b64)
        logo_stream = io.BytesIO(logo_bytes)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_stream, width=Inches(2.0))

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(f"Maquettes {client_name}")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = _hex_to_rgb(primary_color)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(f"Inspirées de {source_url}")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(datetime.now().strftime("%d/%m/%Y"))
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    sep_para = doc.add_paragraph()
    sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = sep_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), primary_color.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.add_paragraph()
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run("Sommaire des maquettes")
    summary_run.bold = True
    summary_run.font.size = Pt(13)
    summary_run.font.color.rgb = _hex_to_rgb(primary_color)

    for i, m in enumerate(mockups, 1):
        item = doc.add_paragraph(style="List Bullet")
        item_run = item.add_run(f"{i}. {m['title']}")
        item_run.font.size = Pt(11)

    # ── Une page par maquette ─────────────────────────────────────────────────
    for i, mockup in enumerate(mockups, 1):
        doc.add_page_break()

        # Titre écran — espacement zéro
        header_para = doc.add_paragraph()
        _zero_spacing(header_para)
        header_run = header_para.add_run(f"Écran {i}/{len(mockups)} — {mockup['title']}")
        header_run.bold = True
        header_run.font.size = Pt(13)
        header_run.font.color.rgb = _hex_to_rgb(primary_color)

        pPr = header_para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), secondary_color.lstrip("#"))
        pBdr.append(bottom)
        pPr.append(pBdr)

        # URL source — espacement zéro, juste en dessous du titre
        url_para = doc.add_paragraph()
        _zero_spacing(url_para)
        url_run = url_para.add_run(f"Source : {mockup.get('url', source_url)}")
        url_run.font.size = Pt(8)
        url_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        url_run.italic = True

        # Image — directement après l'URL, sans paragraphe vide
        try:
            screenshot_bytes = await html_to_screenshot(mockup["html"])
            img_stream = io.BytesIO(screenshot_bytes)
            img_para = doc.add_paragraph()
            _zero_spacing(img_para)
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            # Contraindre largeur ET hauteur max pour rester dans la page
            run.add_picture(
                img_stream,
                width=Inches(img_width_inches),
                height=Inches(img_height_inches),
            )
        except Exception as e:
            err_para = doc.add_paragraph()
            err_run = err_para.add_run(f"⚠ Impossible de rendre la maquette : {e}")
            err_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        # Note bas de page — espacement zéro
        note_para = doc.add_paragraph()
        _zero_spacing(note_para)
        note_run = note_para.add_run(
            f"★ Maquette générée automatiquement — inspirée de {source_url} "
            f"et rebrandée aux couleurs {client_name}"
        )
        note_run.font.size = Pt(7)
        note_run.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
        note_run.italic = True

    # ── Sérialisation ─────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()