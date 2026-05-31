"""
sfd_methods.py — Pipeline unifié de génération SFD.
"""

import os
import json
import base64
import asyncio
from datetime import datetime
from typing import List, Optional, Callable, Awaitable

from fastapi import UploadFile
from google import genai
from google.genai import types

from schemas.sfd_schema import (
    SFD, WebsiteAnalysis, ActeurSFD, ModuleSFD,
    SerieStatistique, CasUtilisation, ApiEndpoint, ExigenceNonFonctionnelle
)
from methods.web_explorer import explore_website

# ============================================================================
# CONFIG
# ============================================================================

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
GEMINI_MODEL   = "gemini-2.5-flash"
OUTPUT_DIR     = "./tmp/outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

ProgressCallback = Callable[[str, str], Awaitable[None]]


# ============================================================================
# EXTRACTION CONTENU FICHIERS
# ============================================================================

async def extract_files_content(files: List[UploadFile]) -> list:
    if not files:
        return []

    contents = []

    for uploaded_file in files:
        if not uploaded_file or not uploaded_file.filename:
            continue

        file_bytes   = await uploaded_file.read()
        filename     = uploaded_file.filename.strip()
        content_type = uploaded_file.content_type or ""
        ext          = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

        if ext == "pdf" or "pdf" in content_type:
            contents.append({
                "type": "pdf", "filename": filename,
                "b64": base64.b64encode(file_bytes).decode("utf-8"),
                "mime": "application/pdf"
            })

        elif ext in ("jpg", "jpeg", "png", "webp", "gif") or "image" in content_type:
            mime = content_type if "image" in content_type else f"image/{ext}"
            contents.append({
                "type": "image", "filename": filename,
                "b64": base64.b64encode(file_bytes).decode("utf-8"),
                "mime": mime
            })

        elif ext in ("docx", "doc") or "word" in content_type:
            try:
                import docx
                from io import BytesIO
                doc  = docx.Document(BytesIO(file_bytes))
                text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                contents.append({"type": "text", "filename": filename, "text": text})
            except Exception as e:
                contents.append({"type": "text", "filename": filename, "text": f"[Erreur DOCX: {e}]"})

        elif ext == "txt" or "text/plain" in content_type:
            contents.append({
                "type": "text", "filename": filename,
                "text": file_bytes.decode("utf-8", errors="ignore")
            })

        else:
            try:
                contents.append({
                    "type": "text", "filename": filename,
                    "text": file_bytes.decode("utf-8", errors="ignore")[:5000]
                })
            except Exception:
                contents.append({"type": "text", "filename": filename, "text": "[Format non supporté]"})

    return contents


# ============================================================================
# PROMPT GEMINI
# ============================================================================

def build_sfd_prompt(
    exploration_result: dict,
    client_name: str,
    project_description: Optional[str] = None
) -> str:
    visited = exploration_result.get("visited_urls", [])
    steps   = exploration_result.get("steps_completed", 0)
    history = exploration_result.get("history", [])

    fonctionnement_lines = []
    for h in history:
        action     = h.get("action", "")
        result     = h.get("result", "")
        url_before = h.get("url_before", "")
        url_after  = h.get("url_after", "")

        if url_before != url_after:
            fonctionnement_lines.append(
                f"  → Navigation : {url_before} ➜ {url_after} | Action: {action}"
            )
        else:
            fonctionnement_lines.append(
                f"  → Interaction : {action} | Résultat: {result[:80]}"
            )

    fonctionnement_text = "\n".join(fonctionnement_lines) if fonctionnement_lines else "Aucune exploration disponible."
    pages_visitees_text = "\n".join([f"  - {u}" for u in visited[:30]])
    client_context      = f'"{client_name}"' if client_name else "la cible (à déduire des documents)"
    desc_block          = f"\nDescription client :\n{project_description}\n" if project_description else ""

    return f"""
Tu es un expert senior en spécification fonctionnelle pour des systèmes d'information.

Ta mission : Générer un SFD (Spécification Fonctionnelle Détaillée) COMPLET et PROFESSIONNEL
pour {client_context}.

{desc_block}

══════════════════════════════════════════════════
ANALYSE DU SITE DE RÉFÉRENCE — FONCTIONNEMENT
URL : {exploration_result.get('url', '')}
Titre : {exploration_result.get('title', '')}
Exploration : {steps} actions effectuées
══════════════════════════════════════════════════

PAGES ET SECTIONS DÉCOUVERTES ({len(visited)}) :
{pages_visitees_text}

FONCTIONNEMENT OBSERVÉ (navigations et interactions) :
{fonctionnement_text}

══════════════════════════════════════════════════

IMPORTANT : Base le SFD sur le FONCTIONNEMENT observé :
- Quelles sections existent et comment on y accède
- Quelles interactions sont disponibles (recherche, filtres, formulaires, navigation)
- Comment les données sont organisées et consultées
- Quels parcours utilisateur ont été observés
- Quelles fonctionnalités ont été déclenchées lors de la navigation

NE PAS reproduire le contenu textuel des pages — analyser leur RÔLE FONCTIONNEL.

Les documents fournis donnent le contexte client. Adapte le SFD à ce contexte.

INSTRUCTIONS :
1. Génère minimum 6 modules, 6 cas d'utilisation, 6 endpoints API.
2. Chaque module doit correspondre à une section/fonctionnalité RÉELLEMENT observée.
3. Les cas d'utilisation décrivent des PARCOURS UTILISATEUR concrets.
4. Les endpoints reflètent les interactions et données observées.

FORMAT : JSON valide uniquement, aucun texte avant/après, aucune balise markdown.

{{
  "nom_projet": "string",
  "nom_cible": "string",
  "version": "1.0",
  "date_creation": "{datetime.now().strftime('%Y-%m-%d')}",
  "auteur": "SFD Generator AI",
  "statut": "Draft",
  "contexte_projet": "string (2-3 paragraphes sur le contexte fonctionnel observé)",
  "objectifs": ["string"],
  "perimetre": "string",
  "public_cible": ["string"],
  "analyse_site_reference": {{
    "url_analysee": "{exploration_result.get('url', '')}",
    "titre_site": "{exploration_result.get('title', '')}",
    "fonctionnalites_identifiees": ["string — fonctionnalité observée"],
    "themes_couverts": ["string — section ou domaine fonctionnel"],
    "structure_navigation": ["string — chemin de navigation observé"],
    "technologies_detectees": ["string"],
    "recommandations": ["string"],
    "pages_visitees": {json.dumps(visited[:15])},
    "actions_effectuees": {steps}
  }},
  "acteurs": [
    {{"nom": "string", "role": "string", "droits": ["string"]}}
  ],
  "modules": [
    {{
      "nom": "string",
      "description": "string — rôle fonctionnel du module",
      "fonctionnalites": ["string — fonctionnalité précise et concrète"],
      "source_inspiration": "string — URL ou section observée"
    }}
  ],
  "series_statistiques": [
    {{
      "categorie": "string",
      "description": "string",
      "format_echange": "string",
      "frequence_mise_a_jour": "string"
    }}
  ],
  "cas_utilisation": [
    {{
      "identifiant": "UC-001",
      "titre": "string",
      "acteur": "string",
      "description": "string",
      "preconditions": ["string"],
      "etapes": ["1. string", "2. string", "3. string"],
      "postconditions": ["string"]
    }}
  ],
  "api_endpoints": [
    {{
      "methode": "GET",
      "chemin": "/api/v1/...",
      "description": "string",
      "parametres": ["string"]
    }}
  ],
  "exigences_non_fonctionnelles": [
    {{
      "categorie": "string",
      "description": "string",
      "critere_acceptance": "string mesurable"
    }}
  ],
  "architecture_technique": "string",
  "contraintes": ["string"],
  "glossaire": {{"terme": "définition"}}
}}
"""


# ============================================================================
# GÉNÉRATION GEMINI
# ============================================================================

async def generate_sfd_with_gemini(
    files_content: list,
    exploration_result: dict,
    client_name: str,
    project_description: Optional[str] = None
) -> SFD:
    if not GOOGLE_API_KEY:
        raise ValueError("GOOGLE_API_KEY non configurée dans .env")

    client = genai.Client(api_key=GOOGLE_API_KEY)

    parts = [types.Part.from_text(text=build_sfd_prompt(
        exploration_result, client_name, project_description
    ))]

    for item in files_content:
        if item["type"] in ("pdf", "image"):
            parts.append(types.Part.from_bytes(
                data=base64.b64decode(item["b64"]),
                mime_type=item["mime"]
            ))
            parts.append(types.Part.from_text(text=f"[Document: {item['filename']}]"))
        elif item["type"] == "text":
            parts.append(types.Part.from_text(
                text=f"\n--- Document: {item['filename']} ---\n{item['text'][:3000]}\n"
            ))

    response = await asyncio.to_thread(
        client.models.generate_content,
        model=GEMINI_MODEL,
        contents=parts,
        config=types.GenerateContentConfig(
            temperature=0.3,
            max_output_tokens=65536,
        )
    )

    raw_text = response.text.strip()

    if raw_text.startswith("```"):
        lines = raw_text.split("\n")
        end   = len(lines) - 1 if lines[-1].strip() == "```" else len(lines)
        raw_text = "\n".join(lines[1:end])

    sfd_dict = None
    try:
        sfd_dict = json.loads(raw_text)
    except json.JSONDecodeError:
        print("⚠️ JSON tronqué détecté — tentative de réparation...")
        for i in range(len(raw_text), 0, -100):
            try:
                sfd_dict = json.loads(raw_text[:i])
                print(f"✅ JSON réparé à {i} caractères")
                break
            except json.JSONDecodeError:
                continue

    if sfd_dict is None:
        raise ValueError(
            f"Impossible de parser le JSON Gemini.\n"
            f"Début : {raw_text[:500]}\nFin : {raw_text[-200:]}"
        )

    return SFD(**sfd_dict)


# ============================================================================
# GÉNÉRATION WORD
# ============================================================================

def save_sfd_as_word(sfd: SFD, nom_fichier: str) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    C1 = RGBColor(0, 51, 102)
    C2 = RGBColor(0, 102, 153)
    C3 = RGBColor(0, 128, 96)

    def h(text, level=1):
        heading = doc.add_heading(text, level=level)
        color = C1 if level == 1 else (C2 if level == 2 else C3)
        for run in heading.runs:
            run.font.color.rgb = color

    def bullet(text):
        doc.add_paragraph(str(text), style='List Bullet')

    def table_block(headers, rows):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = 'Table Grid'
        for i, header in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = header
            for run in cell.paragraphs[0].runs:
                run.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = C1
        for row in rows:
            r = t.add_row()
            for i, val in enumerate(row):
                r.cells[i].text = str(val or "—")
        doc.add_paragraph()

    # Page de garde
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SPÉCIFICATION FONCTIONNELLE DÉTAILLÉE")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = C1

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(sfd.nom_projet)
    r2.bold = True
    r2.font.size = Pt(16)

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(
        f"Cible : {sfd.nom_cible}\n"
        f"Version : {sfd.version}   |   Date : {sfd.date_creation}   |   Statut : {sfd.statut}"
    )
    doc.add_page_break()

    h("1. Contexte et Objectifs")
    doc.add_paragraph(sfd.contexte_projet)
    h("1.1 Objectifs", 2)
    for o in sfd.objectifs:
        bullet(o)
    h("1.2 Périmètre", 2)
    doc.add_paragraph(sfd.perimetre)
    h("1.3 Public cible", 2)
    for pb in sfd.public_cible:
        bullet(pb)

    if sfd.analyse_site_reference:
        ref = sfd.analyse_site_reference
        h("2. Analyse Fonctionnelle du Site de Référence")
        p = doc.add_paragraph()
        p.add_run("URL analysée : ").bold = True
        p.add_run(ref.url_analysee or "—")
        p3 = doc.add_paragraph()
        p3.add_run(f"Exploration : {ref.actions_effectuees} actions, {len(ref.pages_visitees)} pages visitées.")
        if ref.fonctionnalites_identifiees:
            h("2.1 Fonctionnalités observées", 2)
            for f in ref.fonctionnalites_identifiees:
                bullet(f)
        if ref.structure_navigation:
            h("2.2 Structure de navigation", 2)
            for s in ref.structure_navigation:
                bullet(s)
        if ref.themes_couverts:
            h("2.3 Domaines fonctionnels", 2)
            for t in ref.themes_couverts:
                bullet(t)
        if ref.technologies_detectees:
            h("2.4 Technologies détectées", 2)
            for tech in ref.technologies_detectees:
                bullet(tech)
        if ref.recommandations:
            h("2.5 Recommandations", 2)
            for r in ref.recommandations:
                bullet(r)
        if ref.pages_visitees:
            h("2.6 Pages visitées", 2)
            for page in ref.pages_visitees[:20]:
                bullet(page)

    if sfd.acteurs:
        h("3. Acteurs du Système")
        table_block(
            ["Acteur", "Rôle", "Droits"],
            [[a.nom, a.role, "\n".join(a.droits)] for a in sfd.acteurs]
        )

    h("4. Modules Fonctionnels")
    for i, m in enumerate(sfd.modules, 1):
        h(f"4.{i} {m.nom}", 2)
        doc.add_paragraph(m.description)
        if m.source_inspiration:
            p = doc.add_paragraph()
            r = p.add_run(f"Source : {m.source_inspiration}")
            r.italic = True
            r.font.color.rgb = C2
        if m.fonctionnalites:
            h("Fonctionnalités :", 3)
            for f in m.fonctionnalites:
                bullet(f)

    if sfd.series_statistiques:
        h("5. Données et Séries")
        table_block(
            ["Catégorie", "Description", "Format", "Fréquence"],
            [[s.categorie, s.description, s.format_echange or "—", s.frequence_mise_a_jour or "—"]
             for s in sfd.series_statistiques]
        )

    if sfd.cas_utilisation:
        h("6. Cas d'Utilisation")
        for uc in sfd.cas_utilisation:
            h(f"{uc.identifiant} — {uc.titre}", 2)
            p = doc.add_paragraph()
            p.add_run("Acteur : ").bold = True
            p.add_run(uc.acteur)
            doc.add_paragraph(uc.description)
            if uc.preconditions:
                h("Préconditions :", 3)
                for pr in uc.preconditions:
                    bullet(pr)
            if uc.etapes:
                h("Étapes :", 3)
                for j, e in enumerate(uc.etapes, 1):
                    doc.add_paragraph(f"{j}. {e}")
            if uc.postconditions:
                h("Postconditions :", 3)
                for pc in uc.postconditions:
                    bullet(pc)

    if sfd.api_endpoints:
        h("7. API")
        table_block(
            ["Méthode", "Endpoint", "Description", "Paramètres"],
            [[ep.methode, ep.chemin, ep.description, ", ".join(ep.parametres or [])]
             for ep in sfd.api_endpoints]
        )

    if sfd.exigences_non_fonctionnelles:
        h("8. Exigences Non Fonctionnelles")
        for enf in sfd.exigences_non_fonctionnelles:
            h(enf.categorie, 2)
            doc.add_paragraph(enf.description)
            if enf.critere_acceptance:
                p = doc.add_paragraph()
                p.add_run("Critère : ").bold = True
                p.add_run(enf.critere_acceptance)

    if sfd.architecture_technique:
        h("9. Architecture Technique")
        doc.add_paragraph(sfd.architecture_technique)

    if sfd.contraintes:
        h("10. Contraintes")
        for c in sfd.contraintes:
            bullet(c)

    if sfd.glossaire:
        h("11. Glossaire")
        table_block(["Terme", "Définition"], [[k, v] for k, v in sfd.glossaire.items()])

    file_path = os.path.join(OUTPUT_DIR, f"{nom_fichier}.docx")
    doc.save(file_path)
    return file_path


# ============================================================================
# PIPELINE PRINCIPAL
# ============================================================================

async def process_sfd(
    target_url: Optional[str],
    files: List[UploadFile],
    client_name: Optional[str] = None,
    project_description: Optional[str] = None,
    max_exploration_calls: int = 6,
    on_progress: Optional[ProgressCallback] = None
) -> tuple:
    if not GOOGLE_API_KEY:
        raise ValueError("GOOGLE_API_KEY non configurée")

    async def notify(stage: str, message: str):
        print(f"[{stage}] {message}")
        if on_progress:
            await on_progress(stage, message)

    effective_name = client_name.strip() if client_name and client_name.strip() else "portail_cible"

    # Étape 1 : Exploration web
    if target_url and target_url.strip().startswith("http"):
        exploration_result = await explore_website(
            target_url=target_url.strip(),
            gemini_api_key=GOOGLE_API_KEY,
            max_gemini_calls=max_exploration_calls,
            actions_per_call=6,
            on_progress=on_progress  # déléguer directement
        )
    else:
        exploration_result = {
            "url": target_url or "", "title": "",
            "history": [], "visited_urls": [],
            "steps_completed": 0,
            "text_summary": "Aucune URL fournie.",
            "screenshots": []
        }

    # Étape 2 : Fichiers
    await notify("analysis", "📄 Analyse des documents fournis...")
    files_content = await extract_files_content(files) if files else []

    # Étape 3 : Génération SFD
    await notify("generation", "🧠 Génération du SFD avec Gemini AI...")
    sfd = await generate_sfd_with_gemini(
        files_content=files_content,
        exploration_result=exploration_result,
        client_name=effective_name,
        project_description=project_description
    )

    # Étape 4 : Word
    await notify("word", "📝 Création du document Word...")
    timestamp  = datetime.now().strftime("%Y%m%d_%H%M%S")
    nom_fichier = f"SFD_{effective_name.replace(' ', '_')}_{timestamp}"
    word_path  = save_sfd_as_word(sfd, nom_fichier)

    exploration_summary = (
        f"Exploration : {exploration_result['steps_completed']} actions, "
        f"{len(exploration_result['visited_urls'])} pages visitées sur {target_url}"
        if target_url else "Aucune exploration web"
    )

    await notify("done", f"✅ SFD prêt — {exploration_summary}")

    return sfd, os.path.basename(word_path), exploration_summary