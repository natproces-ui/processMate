"""
word_methods.py - Génération de documents Word professionnels pour SFD
Mise en forme professionnelle, sans emojis, sans espaces inutiles
"""

import os
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Any, List
from schemas.format1_schema import Format1SFD
from schemas.format2_schema import Format2SFD

OUTPUT_DIR = "./tmp/outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Palette de couleurs
COLOR_TITLE      = RGBColor(0x1F, 0x39, 0x64)
COLOR_H1         = RGBColor(0x1F, 0x39, 0x64)
COLOR_H2         = RGBColor(0x2E, 0x74, 0xB5)
COLOR_H3         = RGBColor(0x2E, 0x74, 0xB5)
COLOR_TABLE_HEAD = "2E74B5"
COLOR_BORDER     = "BFBFBF"


# ─── HELPERS XML ─────────────────────────────────────────────────────────────

def set_cell_background(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, color="BFBFBF"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    pPr.append(spacing)


def add_page_number_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.clear()
    run = fp.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    run2 = fp.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    run2._r.append(instrText)
    run3 = fp.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar2)
    for r in fp.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def add_rule(doc: Document, color="1F3964", thickness=8):
    para = doc.add_paragraph()
    set_spacing(para, before=0, after=60)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─── STYLES ──────────────────────────────────────────────────────────────────

def _run(para, text, bold=False, italic=False, size=10, color=None):
    run = para.add_run(str(text))
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    return run


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=120)
    _run(p, text, bold=True, size=26, color=COLOR_TITLE)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=200)
    _run(p, text, size=16, color=COLOR_H2)


def add_h1(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=240, after=80)
    _run(p, text.upper(), bold=True, size=13, color=COLOR_H1)
    add_rule(doc, color="1F3964", thickness=8)


def add_h2(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=160, after=60)
    _run(p, text, bold=True, size=11, color=COLOR_H2)


def add_h3(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=120, after=40)
    _run(p, text, bold=True, size=10, color=COLOR_H3)


def add_h4(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=80, after=20)
    _run(p, text, bold=True, italic=True, size=10, color=RGBColor(0x40, 0x40, 0x40))


def add_body(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=60)
    _run(p, str(text), size=10)


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    set_spacing(p, before=0, after=30)
    p.clear()
    _run(p, str(text), size=10)


def add_bullet_list(doc, items: List[str]):
    for item in items:
        if item and str(item).strip():
            add_bullet(doc, str(item))


def add_kv(doc, key, value):
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=40)
    _run(p, f"{key}: ", bold=True, size=10)
    _run(p, str(value), size=10, color=RGBColor(0x40, 0x40, 0x40))


# ─── TABLEAUX ────────────────────────────────────────────────────────────────

def add_metadata_table(doc, rows: Dict[str, str]):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for key, value in rows.items():
        row = table.add_row()
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Cm(4.5)
        c1.width = Cm(11)
        set_cell_background(c0, "E8EEF5")
        set_cell_border(c0)
        set_cell_border(c1)
        r0 = c0.paragraphs[0].add_run(str(key))
        r0.font.name = 'Calibri'; r0.font.size = Pt(9); r0.font.bold = True
        r0.font.color.rgb = COLOR_H1
        r1 = c1.paragraphs[0].add_run(str(value))
        r1.font.name = 'Calibri'; r1.font.size = Pt(9)
    doc.add_paragraph()


def add_user_story_block(doc, us):
    # En-tête
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    row = t.rows[0]
    for cell, txt, w in zip(row.cells,
                             [str(us.id), str(us.titre), f"{us.priorite} | {us.estimation}"],
                             [Cm(2), Cm(10), Cm(4)]):
        cell.width = w
        set_cell_background(cell, COLOR_TABLE_HEAD)
        set_cell_border(cell, color=COLOR_TABLE_HEAD)
        r = cell.paragraphs[0].add_run(txt)
        r.font.name = 'Calibri'; r.font.size = Pt(9); r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Corps
    bt = doc.add_table(rows=1, cols=1)
    bt.alignment = WD_TABLE_ALIGNMENT.LEFT
    bc = bt.rows[0].cells[0]
    bc.width = Cm(16)
    set_cell_background(bc, "F7FAFD")
    set_cell_border(bc)

    # Formulation
    p = bc.add_paragraph()
    for label, val in [("En tant que ", us.en_tant_que + ", "),
                        ("je veux ", us.je_veux + " "),
                        ("afin de ", us.afin_de + ".")]:
        r = p.add_run(label); r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(10)
        r = p.add_run(val);   r.font.name = 'Calibri'; r.font.size = Pt(10)

    ps = bc.add_paragraph()
    rs = ps.add_run(f"Statut : {us.statut}")
    rs.font.name = 'Calibri'; rs.font.size = Pt(9); rs.font.italic = True
    rs.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    if us.criteres_acceptation:
        pc = bc.add_paragraph()
        rc = pc.add_run("Criteres d'acceptation :")
        rc.font.name = 'Calibri'; rc.font.size = Pt(10); rc.font.bold = True
        for c in us.criteres_acceptation:
            p2 = bc.add_paragraph(f"  - {str(c)}")
            if p2.runs: p2.runs[0].font.name = 'Calibri'; p2.runs[0].font.size = Pt(10)

    if us.regles_metier:
        pr = bc.add_paragraph()
        rr = pr.add_run("Regles metier :")
        rr.font.name = 'Calibri'; rr.font.size = Pt(10); rr.font.bold = True
        for r_ in us.regles_metier:
            p2 = bc.add_paragraph(f"  - {str(r_)}")
            if p2.runs: p2.runs[0].font.name = 'Calibri'; p2.runs[0].font.size = Pt(10)

    sp = doc.add_paragraph()
    set_spacing(sp, before=0, after=60)


# ─── PAGE DE GARDE ───────────────────────────────────────────────────────────

def add_cover_page(doc, title, subtitle, metadata):
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=400)
    add_title(doc, title)
    add_rule(doc, color="2E74B5", thickness=16)
    add_subtitle(doc, subtitle)
    doc.add_paragraph()
    add_metadata_table(doc, metadata)
    doc.add_page_break()


# ─── FORMAT 1 ────────────────────────────────────────────────────────────────

def build_format1_doc(sfd: Format1SFD) -> Document:
    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.5)
    add_page_number_footer(doc)

    add_cover_page(doc,
        "Specifications Fonctionnelles Detaillees",
        sfd.nom_projet,
        {"Version": str(sfd.version), "Date": str(sfd.date), "Auteur": str(sfd.auteur or "N/A")}
    )

    add_h1(doc, "1. Contexte et Perimetre")
    add_h2(doc, "1.1 Presentation du projet")
    add_body(doc, sfd.contexte.presentation)
    add_h2(doc, "1.2 Objectifs metier")
    add_bullet_list(doc, sfd.contexte.objectifs_metier)
    add_h2(doc, "1.3 Perimetre fonctionnel")
    add_body(doc, sfd.contexte.perimetre)
    add_h2(doc, "1.4 Acteurs et roles")
    for a in sfd.contexte.acteurs:
        role = a.get('role', '') if isinstance(a, dict) else a.role
        desc = a.get('description', '') if isinstance(a, dict) else a.description
        add_bullet(doc, f"{role} : {desc}")

    doc.add_page_break()

    add_h1(doc, "2. Description Generale")
    add_h2(doc, "2.1 Architecture fonctionnelle")
    add_body(doc, sfd.description_generale.architecture_fonctionnelle)
    add_h2(doc, "2.2 Flux principaux")
    add_bullet_list(doc, sfd.description_generale.flux_principaux)
    add_h2(doc, "2.3 Regles de gestion globales")
    add_bullet_list(doc, sfd.description_generale.regles_gestion)

    doc.add_page_break()

    add_h1(doc, "3. Exigences Fonctionnelles Detaillees")
    for i, mod in enumerate(sfd.modules, 1):
        add_h2(doc, f"3.{i}  {mod.nom}  [{mod.id}]")
        add_body(doc, mod.description)
        for j, fn in enumerate(mod.fonctions, 1):
            add_h3(doc, f"3.{i}.{j}  {fn.nom}  [{fn.id}]")
            add_body(doc, fn.description)
            if fn.regles_metier:
                add_h4(doc, "Regles metier"); add_bullet_list(doc, fn.regles_metier)
            if fn.donnees_entree:
                add_h4(doc, "Donnees en entree"); add_bullet_list(doc, fn.donnees_entree)
            if fn.donnees_sortie:
                add_h4(doc, "Donnees en sortie"); add_bullet_list(doc, fn.donnees_sortie)
            if fn.cas_nominal:
                add_h4(doc, "Cas nominal"); add_body(doc, fn.cas_nominal)
            if fn.cas_erreur:
                add_h4(doc, "Cas d'erreur"); add_bullet_list(doc, fn.cas_erreur)
            if fn.contraintes:
                add_h4(doc, "Contraintes"); add_bullet_list(doc, fn.contraintes)

    doc.add_page_break()

    add_h1(doc, "4. Exigences Non Fonctionnelles")
    enf = sfd.exigences_non_fonctionnelles
    if enf.performance:    add_h2(doc, "4.1 Performance");    add_bullet_list(doc, enf.performance)
    if enf.securite:       add_h2(doc, "4.2 Securite");       add_bullet_list(doc, enf.securite)
    if enf.disponibilite:  add_h2(doc, "4.3 Disponibilite");  add_bullet_list(doc, enf.disponibilite)
    if enf.scalabilite:    add_h2(doc, "4.4 Scalabilite");    add_bullet_list(doc, enf.scalabilite)

    doc.add_page_break()

    add_h1(doc, "5. Contraintes Techniques")
    ct = sfd.contraintes_techniques
    if ct.environnement:  add_h2(doc, "5.1 Environnement");          add_bullet_list(doc, ct.environnement)
    if ct.technologies:   add_h2(doc, "5.2 Technologies");            add_bullet_list(doc, ct.technologies)
    if ct.normes:         add_h2(doc, "5.3 Normes et conformite");    add_bullet_list(doc, ct.normes)

    if sfd.glossaire or sfd.notes:
        doc.add_page_break()
        add_h1(doc, "6. Annexes")
        if sfd.glossaire:
            add_h2(doc, "6.1 Glossaire")
            for t, d in sfd.glossaire.items():
                add_kv(doc, str(t), str(d))
        if sfd.notes:
            add_h2(doc, "6.2 Notes complementaires")
            add_body(doc, sfd.notes)

    return doc


# ─── FORMAT 2 ────────────────────────────────────────────────────────────────

def build_format2_doc(sfd: Format2SFD) -> Document:
    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.5)
    add_page_number_footer(doc)

    add_cover_page(doc,
        "Specifications Fonctionnelles Agiles",
        sfd.nom_projet,
        {"Version": str(sfd.version), "Date": str(sfd.date),
         "Product Owner": str(sfd.product_owner or "N/A"),
         "Scrum Master": str(sfd.scrum_master or "N/A")}
    )

    add_h1(doc, "1. Vision Produit")
    add_h2(doc, "1.1 Probleme a resoudre")
    add_body(doc, sfd.vision_produit.probleme)
    add_h2(doc, "1.2 Utilisateurs cibles")
    add_bullet_list(doc, sfd.vision_produit.utilisateurs_cibles)
    add_h2(doc, "1.3 Valeur apportee")
    add_body(doc, sfd.vision_produit.valeur_apportee)
    add_h2(doc, "1.4 Objectifs")
    add_bullet_list(doc, sfd.vision_produit.objectifs)

    doc.add_page_break()

    add_h1(doc, "2. Epics et User Stories")
    for i, epic in enumerate(sfd.epics, 1):
        add_h2(doc, f"Epic {i} : {epic.nom}  [{epic.id}]")
        add_body(doc, epic.description)
        add_kv(doc, "Objectif", epic.objectif)
        for us in epic.user_stories:
            add_user_story_block(doc, us)

    doc.add_page_break()

    if sfd.regles_metier:
        add_h1(doc, "3. Regles Metier Transverses")
        for regle in sfd.regles_metier:
            add_h2(doc, f"{regle.nom}  [{regle.id}]")
            add_body(doc, regle.description)
            if regle.impact:
                add_kv(doc, "Impact", ", ".join(str(x) for x in regle.impact))
        doc.add_page_break()

    add_h1(doc, "4. Modele de Donnees")
    if sfd.modele_data.entites:
        add_h2(doc, "4.1 Entites principales")
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = table.rows[0]
        for cell, txt in zip(hdr.cells, ["Entite", "Attributs"]):
            set_cell_background(cell, COLOR_TABLE_HEAD)
            set_cell_border(cell, color=COLOR_TABLE_HEAD)
            r = cell.paragraphs[0].add_run(txt)
            r.font.name = 'Calibri'; r.font.size = Pt(9); r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for idx, entite in enumerate(sfd.modele_data.entites):
            nom = entite.get('nom', '') if isinstance(entite, dict) else entite.nom
            attributs = entite.get('attributs', '') if isinstance(entite, dict) else entite.attributs
            row = table.add_row()
            bg = "EAF0F8" if idx % 2 == 0 else "FFFFFF"
            for cell, txt in zip(row.cells, [str(nom), str(attributs)]):
                set_cell_background(cell, bg)
                set_cell_border(cell)
                r = cell.paragraphs[0].add_run(txt)
                r.font.name = 'Calibri'; r.font.size = Pt(9)
        doc.add_paragraph()

    if sfd.modele_data.relations:
        add_h2(doc, "4.2 Relations")
        add_bullet_list(doc, sfd.modele_data.relations)

    doc.add_page_break()

    if sfd.workflows:
        add_h1(doc, "5. Workflows")
        for wf in sfd.workflows:
            add_h2(doc, wf.nom)
            if wf.description:
                add_body(doc, wf.description)
            for k, etape in enumerate(wf.etapes, 1):
                add_bullet(doc, f"Etape {k} : {str(etape)}")
        doc.add_page_break()

    if sfd.definition_of_done:
        add_h1(doc, "6. Definition of Done")
        add_bullet_list(doc, sfd.definition_of_done)
        doc.add_page_break()

    if sfd.notes:
        add_h1(doc, "7. Notes Complementaires")
        add_body(doc, sfd.notes)

    return doc


# ─── API PUBLIQUE ─────────────────────────────────────────────────────────────

def generate_word_document(sfd_data: Dict[str, Any], format_type: str, nom_fichier: str) -> str:
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, f"{nom_fichier}.docx")
    if format_type == "format1":
        doc = build_format1_doc(Format1SFD(**sfd_data))
    elif format_type == "format2":
        doc = build_format2_doc(Format2SFD(**sfd_data))
    else:
        raise ValueError(f"Format non supporte : {format_type}")
    doc.save(output_path)
    return output_path


def generate_word_bytes(sfd_data: Dict[str, Any], format_type: str) -> bytes:
    if format_type == "format1":
        doc = build_format1_doc(Format1SFD(**sfd_data))
    elif format_type == "format2":
        doc = build_format2_doc(Format2SFD(**sfd_data))
    else:
        raise ValueError(f"Format non supporte : {format_type}")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()