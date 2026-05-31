"""
word_renderer.py — Génération du document Word SFD avec support multi-thèmes et sommaire TOC.
Même structure que html_renderer.py. Thème injecté via SFDTheme (themes.py).
"""

import io
import base64
import urllib.request
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from schemas.schema import SFDDocument
from themes import SFDTheme, get_theme, DEFAULT_THEME
from methods.html_renderer import fix_mermaid

FONT = "Calibri"


# ─── HELPERS XML ──────────────────────────────────────────────────────────────

def _spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"), str(after))
    pPr.append(sp)


def _run(para, text, bold=False, italic=False, size=10, color=None, theme: SFDTheme = None):
    run = para.add_run(str(text))
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    elif theme:
        run.font.color.rgb = theme.color_body
    else:
        run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    return run


def _cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_border(cell, color: str = "BFBFBF"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        borders.append(b)
    tcPr.append(borders)


def _rule(doc, theme: SFDTheme):
    p = doc.add_paragraph()
    _spacing(p, before=0, after=60)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(theme.rule_thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), theme.rule_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _page_number_footer(doc, theme: SFDTheme):
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.clear()
    run = fp.add_run()
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "begin")
    run._r.append(fc)
    run2 = fp.add_run()
    instr = OxmlElement("w:instrText")
    instr.text = " PAGE "
    run2._r.append(instr)
    run3 = fp.add_run()
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    run3._r.append(fc2)
    for r in fp.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ─── SOMMAIRE TOC NATIF ───────────────────────────────────────────────────────

def _add_toc(doc: Document, theme: SFDTheme):
    """
    Insère un sommaire TOC natif Word via champ TOC.
    Word met à jour les numéros de page à l'ouverture du document.
    Structure : titre "Sommaire" + champ TOC1-TOC3.
    """
    # Titre du sommaire avec bordure basse directe
    p_title = doc.add_paragraph()
    _spacing(p_title, before=0, after=160)
    _run(p_title, "SOMMAIRE", bold=True, size=16, color=theme.color_h1)
    pPr = p_title._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(theme.rule_thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), theme.rule_color)
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Paragraphe contenant le champ TOC
    p_toc = doc.add_paragraph()
    _spacing(p_toc, before=120, after=120)

    # Construction du champ TOC via XML OpenXML
    # \o "1-3" = niveaux 1 à 3, \h = hyperliens, \z = masquer onglets, \u = utiliser styles Heading
    run = p_toc.add_run()

    # fldChar begin
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")   # force recalcul à l'ouverture
    run._r.append(fld_begin)

    # instrText : instruction TOC
    run2 = p_toc.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instr)

    # fldChar separate (placeholder texte)
    run3 = p_toc.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_sep)

    # Texte placeholder (remplacé par Word au premier refresh)
    run4 = p_toc.add_run()
    run4.font.size = Pt(10)
    run4.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    run4.font.italic = True
    run4.text = "[ Ouvrez dans Microsoft Word et appuyez sur F9 pour mettre à jour le sommaire ]"

    # fldChar end
    run5 = p_toc.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_end)

    doc.add_page_break()


# ─── STYLES TYPOGRAPHIQUES ────────────────────────────────────────────────────

def _title(doc, text, theme: SFDTheme):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, before=0, after=120)
    _run(p, text, bold=True, size=26, color=theme.color_title, theme=theme)


def _subtitle(doc, text, theme: SFDTheme):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, before=0, after=80)
    _run(p, text, size=14, color=theme.color_h2, theme=theme)


def _h1(doc, text, theme: SFDTheme):
    """H1 avec style Heading 1 pour que le TOC le détecte. Bordure basse directe."""
    p = doc.add_heading(text.upper(), level=1)
    _spacing(p, before=240, after=80)
    for run in p.runs:
        run.font.name = FONT
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = theme.color_h1
    # Bordure basse directement sur le para heading (pas de para séparé)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(theme.rule_thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), theme.rule_color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _h2(doc, text, theme: SFDTheme):
    """H2 avec style Heading 2 pour le TOC."""
    p = doc.add_heading(text, level=2)
    _spacing(p, before=160, after=60)
    for run in p.runs:
        run.font.name = FONT
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = theme.color_h2
    return p


def _h3(doc, text, theme: SFDTheme):
    """H3 avec style Heading 3 pour le TOC."""
    p = doc.add_heading(text, level=3)
    _spacing(p, before=120, after=40)
    for run in p.runs:
        run.font.name = FONT
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = theme.color_h3
    return p


def _body(doc, text, theme: SFDTheme):
    p = doc.add_paragraph()
    _spacing(p, before=0, after=60)
    _run(p, str(text), size=10, theme=theme)


def _bullet(doc, text, theme: SFDTheme):
    p = doc.add_paragraph(style="List Bullet")
    _spacing(p, before=0, after=30)
    p.clear()
    _run(p, str(text), size=10, theme=theme)


def _kv(doc, key, value, theme: SFDTheme):
    p = doc.add_paragraph()
    _spacing(p, before=0, after=40)
    _run(p, f"{key}: ", bold=True, size=10, color=theme.color_h2, theme=theme)
    _run(p, str(value), size=10, theme=theme)


def _separator(doc):
    p = doc.add_paragraph()
    _spacing(p, before=120, after=120)


# ─── TABLEAUX ─────────────────────────────────────────────────────────────────

def _add_table(doc, headers: list, rows: list, theme: SFDTheme):
    n_cols = len(headers)
    table = doc.add_table(rows=1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        _cell_bg(cell, theme.table_header_bg)
        _cell_border(cell, color=theme.table_header_bg)
        r = cell.paragraphs[0].add_run(str(h))
        r.font.name = FONT
        r.font.size = Pt(10)
        r.font.bold = True
        # Texte en-tête : blanc ou sombre selon le thème
        fg = theme.table_header_fg
        r.font.color.rgb = RGBColor(
            int(fg[0:2], 16), int(fg[2:4], 16), int(fg[4:6], 16)
        )

    for idx, row_data in enumerate(rows):
        row = table.add_row()
        bg = theme.table_row_alt if idx % 2 == 1 else theme.table_row_normal
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            _cell_bg(cell, bg)
            _cell_border(cell, color=theme.table_border)
            r = cell.paragraphs[0].add_run(str(val))
            r.font.name = FONT
            r.font.size = Pt(10)
            r.font.color.rgb = theme.color_body

    doc.add_paragraph()


def _meta_table(doc, rows: dict, theme: SFDTheme):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for key, val in rows.items():
        row = table.add_row()
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Cm(4)
        _cell_bg(c0, theme.meta_bg)
        _cell_border(c0, color=theme.table_border)
        _cell_border(c1, color=theme.table_border)
        r0 = c0.paragraphs[0].add_run(str(key))
        r0.font.name = FONT
        r0.font.size = Pt(9)
        r0.font.bold = True
        r0.font.color.rgb = theme.meta_fg_key
        r1 = c1.paragraphs[0].add_run(str(val))
        r1.font.name = FONT
        r1.font.size = Pt(9)
        r1.font.color.rgb = theme.meta_fg_val
    doc.add_paragraph()


# ─── SECTIONS ─────────────────────────────────────────────────────────────────

def _write_cover(doc, sfd: SFDDocument, theme: SFDTheme):
    m = sfd.meta
    doc.add_paragraph()
    _subtitle(doc, "SPÉCIFICATION FONCTIONNELLE DÉTAILLÉE", theme)
    _title(doc, m.nom_projet, theme)
    if m.client:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p, before=0, after=200)
        _run(p, m.client, size=14, color=theme.color_h2, theme=theme)

    meta = {"Client": m.client, "Version": m.version,
            "Date": m.date, "Statut": m.statut}
    if m.auteurs:
        meta["Auteurs"] = ", ".join(m.auteurs)

    _meta_table(doc, meta, theme)

    # Bandeau thème en bas de page de garde
    p_band = doc.add_paragraph()
    _spacing(p_band, before=200, after=0)
    pPr = p_band._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "bottom"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "12")
        b.set(qn("w:space"), "1")
        b.set(qn("w:color"), theme.rule_color)
        pBdr.append(b)
    pPr.append(pBdr)
    _run(p_band, f"  {theme.label}  —  Document généré automatiquement",
         italic=True, size=9, color=theme.color_accent, theme=theme)

    doc.add_page_break()

    if m.historique_revisions:
        _h1(doc, "Historique des révisions", theme)
        rows = [[r.version, r.date, r.auteur, r.description]
                for r in m.historique_revisions]
        _add_table(doc, ["Version", "Date", "Auteur", "Description"], rows, theme)

    if sfd.documents_reference:
        _h1(doc, "Documents de référence", theme)
        rows = [[d.nom, d.type, d.version, d.description]
                for d in sfd.documents_reference]
        _add_table(doc, ["Nom", "Type", "Version", "Description"], rows, theme)


def _write_contexte(doc, sfd: SFDDocument, theme: SFDTheme):
    c = sfd.contexte
    _h1(doc, "1. Contexte général", theme)
    if c.presentation_client:
        _h2(doc, "Présentation du client", theme)
        _body(doc, c.presentation_client, theme)
    if c.contexte_projet:
        _h2(doc, "Contexte du projet", theme)
        _body(doc, c.contexte_projet, theme)
    if c.objectifs_metier:
        _h2(doc, "Objectifs métier", theme)
        for obj in c.objectifs_metier:
            _bullet(doc, obj, theme)


def _write_perimetre(doc, sfd: SFDDocument, theme: SFDTheme):
    p = sfd.perimetre
    _h1(doc, "2. Périmètre fonctionnel", theme)
    if p.inclus:
        _h2(doc, "Dans le périmètre", theme)
        for i in p.inclus:
            _bullet(doc, i, theme)
    if p.exclus:
        _h2(doc, "Hors périmètre", theme)
        for i in p.exclus:
            _bullet(doc, i, theme)
    if p.hypotheses:
        _h2(doc, "Hypothèses", theme)
        for i in p.hypotheses:
            _bullet(doc, i, theme)
    if p.contraintes_generales:
        _h2(doc, "Contraintes générales", theme)
        for i in p.contraintes_generales:
            _bullet(doc, i, theme)


def _write_acteurs(doc, sfd: SFDDocument, theme: SFDTheme):
    _h1(doc, "3. Acteurs du système", theme)
    if not sfd.acteurs:
        _body(doc, "Aucun acteur défini.", theme)
        return
    rows = [[a.id, a.nom, a.type, a.role, a.description]
            for a in sfd.acteurs]
    _add_table(doc, ["ID", "Nom", "Type", "Rôle", "Description"], rows, theme)


def _write_cas_utilisation(doc, sfd: SFDDocument, theme: SFDTheme):
    _h1(doc, "4. Cas d'utilisation", theme)
    if not sfd.cas_utilisation:
        _body(doc, "Aucun cas d'utilisation défini.", theme)
        return
    for uc in sfd.cas_utilisation:
        _h2(doc, f"{uc.id} — {uc.nom}", theme)
        _kv(doc, "Acteur principal", uc.acteur_principal, theme)
        if uc.acteurs_secondaires:
            _kv(doc, "Acteurs secondaires", ", ".join(uc.acteurs_secondaires), theme)
        if uc.preconditions:
            _h3(doc, "Préconditions", theme)
            for i in uc.preconditions:
                _bullet(doc, i, theme)
        _h3(doc, "Flux nominal", theme)
        for e in uc.flux_nominal:
            _bullet(doc, f"{e.numero}. {e.description}", theme)
        if uc.flux_alternatifs:
            _h3(doc, "Flux alternatifs", theme)
            for i in uc.flux_alternatifs:
                _bullet(doc, i, theme)
        if uc.flux_erreur:
            _h3(doc, "Gestion des erreurs", theme)
            for i in uc.flux_erreur:
                _bullet(doc, i, theme)
        if uc.postconditions:
            _h3(doc, "Postconditions", theme)
            for i in uc.postconditions:
                _bullet(doc, i, theme)
        _separator(doc)


def _write_modules(doc, sfd: SFDDocument, theme: SFDTheme):
    _h1(doc, "5. Modules et fonctions", theme)
    if not sfd.modules:
        _body(doc, "Aucun module défini.", theme)
        return
    for mod in sfd.modules:
        _h2(doc, f"{mod.id} — {mod.nom}", theme)
        if mod.description:
            _body(doc, mod.description, theme)
        for fn in mod.fonctions:
            _h3(doc, f"{fn.id} — {fn.nom}", theme)
            _kv(doc, "Priorité", fn.priorite, theme)
            if fn.description:
                _body(doc, fn.description, theme)
            if fn.donnees_entree:
                _kv(doc, "Données d'entrée", " | ".join(fn.donnees_entree), theme)
            if fn.donnees_sortie:
                _kv(doc, "Données de sortie", " | ".join(fn.donnees_sortie), theme)
            if fn.regles_gestion_ids:
                _kv(doc, "Règles de gestion", ", ".join(fn.regles_gestion_ids), theme)
            if fn.contraintes:
                _kv(doc, "Contraintes", " | ".join(fn.contraintes), theme)


def _write_regles_gestion(doc, sfd: SFDDocument, theme: SFDTheme):
    _h1(doc, "6. Règles de gestion", theme)
    if not sfd.regles_gestion:
        _body(doc, "Aucune règle définie.", theme)
        return
    rows = [
        [rg.id, rg.type, rg.description, ", ".join(rg.fonctions_concernees)]
        for rg in sfd.regles_gestion
    ]
    _add_table(doc, ["ID", "Type", "Description", "Fonctions concernées"], rows, theme)


def _write_specifications_donnees(doc, sfd: SFDDocument, theme: SFDTheme):
    _h1(doc, "7. Spécifications des données", theme)
    if not sfd.specifications_donnees:
        _body(doc, "Aucune spécification définie.", theme)
        return
    for spec in sfd.specifications_donnees:
        _h2(doc, spec.entite, theme)
        if spec.description:
            _body(doc, spec.description, theme)
        if spec.attributs:
            _add_table(doc, ["Attributs"], [[a] for a in spec.attributs], theme)
        if spec.flux_associes:
            _kv(doc, "Flux associés", ", ".join(spec.flux_associes), theme)


def _write_interfaces(doc, sfd: SFDDocument, theme: SFDTheme):
    _h1(doc, "8. Interfaces", theme)
    if sfd.interfaces.interfaces_ui:
        _h2(doc, "Interfaces utilisateur (IHM)", theme)
        for ihm in sfd.interfaces.interfaces_ui:
            _h3(doc, f"{ihm.id} — {ihm.nom_ecran}", theme)
            if ihm.description:
                _body(doc, ihm.description, theme)
            if ihm.acteur:
                _kv(doc, "Acteur", ihm.acteur, theme)
            if ihm.elements:
                rows = [[e.nom, e.type, e.description, "Oui" if e.obligatoire else "Non"]
                        for e in ihm.elements]
                _add_table(doc, ["Élément", "Type", "Description", "Obligatoire"], rows, theme)

    if sfd.interfaces.interfaces_externes:
        _h2(doc, "Interfaces externes", theme)
        rows = [
            [i.id, i.nom, i.type, i.systeme_tiers, i.description, i.format_echange]
            for i in sfd.interfaces.interfaces_externes
        ]
        _add_table(doc, ["ID", "Nom", "Type", "Système tiers", "Description", "Format"], rows, theme)


def _write_exigences_nf(doc, sfd: SFDDocument, theme: SFDTheme):
    _h1(doc, "9. Exigences non-fonctionnelles", theme)
    enf = sfd.exigences_non_fonctionnelles
    for label, items in [
        ("Performance", enf.performance),
        ("Sécurité", enf.securite),
        ("Disponibilité", enf.disponibilite),
        ("Ergonomie", enf.ergonomie),
        ("Maintenabilité", enf.maintenabilite),
    ]:
        if items:
            _h2(doc, label, theme)
            for i in items:
                _bullet(doc, i, theme)


def _write_matrice_tracabilite(doc, sfd: SFDDocument, theme: SFDTheme):
    _h1(doc, "10. Matrice de traçabilité", theme)
    if not sfd.matrice_tracabilite:
        _body(doc, "Matrice non définie.", theme)
        return
    rows = [
        [m.id_besoin_source, m.description_besoin, ", ".join(m.fonctions_couvrant)]
        for m in sfd.matrice_tracabilite
    ]
    _add_table(doc, ["ID Besoin", "Description du besoin", "Fonctions couvrant"], rows, theme)


def _mermaid_to_png(mermaid_code: str) -> bytes | None:
    try:
        encoded = base64.urlsafe_b64encode(mermaid_code.encode("utf-8")).decode()
        url = f"https://mermaid.ink/img/{encoded}?type=png"
        req = urllib.request.Request(url, headers={"User-Agent": "SFD-Generator/1.0"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            return resp.read()
    except Exception:
        return None


def _write_schemas_conceptuels(doc, sfd: SFDDocument, theme: SFDTheme):
    if not sfd.schemas_conceptuels:
        return
    _h1(doc, "11. Schémas conceptuels", theme)
    for sch in sfd.schemas_conceptuels:
        _h2(doc, f"{sch.id} — {sch.titre}", theme)
        if sch.description:
            _body(doc, sch.description, theme)
        png = _mermaid_to_png(fix_mermaid(sch.mermaid_code))
        if png:
            doc.add_picture(io.BytesIO(png), width=Cm(15))
        else:
            p = doc.add_paragraph()
            _spacing(p, before=40, after=40)
            run = p.add_run(sch.mermaid_code)
            run.font.name = "Courier New"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        _separator(doc)


def _write_glossaire(doc, sfd: SFDDocument, theme: SFDTheme):
    if not sfd.glossaire:
        return
    _h1(doc, "12. Glossaire", theme)
    rows = [[k, v] for k, v in sfd.glossaire.items()]
    _add_table(doc, ["Terme", "Définition"], rows, theme)


# ─── RENDERER PRINCIPAL ───────────────────────────────────────────────────────

def render_docx(sfd: SFDDocument, style: str = DEFAULT_THEME) -> bytes:
    """
    Génère le document Word (.docx) du SFD avec le thème spécifié.

    Args:
        sfd:   L'objet SFDDocument à rendre.
        style: Nom du thème ('al_maghrib' ou 'corporate_blue').

    Returns:
        bytes du fichier DOCX (in-memory, sans écriture disque).
    """
    theme = get_theme(style)

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    _page_number_footer(doc, theme)

    # ── Page de garde (contient déjà son page_break en fin)
    _write_cover(doc, sfd, theme)

    # ── Sommaire TOC natif (contient déjà son page_break en fin)
    _add_toc(doc, theme)

    # ── Contenu
    _write_contexte(doc, sfd, theme)
    _write_perimetre(doc, sfd, theme)
    _write_acteurs(doc, sfd, theme)
    _write_cas_utilisation(doc, sfd, theme)
    _write_modules(doc, sfd, theme)
    _write_regles_gestion(doc, sfd, theme)
    _write_specifications_donnees(doc, sfd, theme)
    _write_interfaces(doc, sfd, theme)
    _write_exigences_nf(doc, sfd, theme)
    _write_matrice_tracabilite(doc, sfd, theme)
    _write_schemas_conceptuels(doc, sfd, theme)
    _write_glossaire(doc, sfd, theme)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()